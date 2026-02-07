"""
DOCX Document Parser
Extracts structured data from Word documents for format checking.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any, Optional
import re
import os


class DocxParser:
    """Parser for Microsoft Word (.docx) documents."""

    # Alignment mapping
    ALIGNMENT_MAP = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    
    # Font name mapping (Chinese to English)
    FONT_NAME_MAP = {
        "宋体": "SimSun",
        "SimSun": "SimSun",
        "黑体": "SimHei",
        "SimHei": "SimHei",
        "微软雅黑": "Microsoft YaHei",
        "Microsoft YaHei": "Microsoft YaHei",
        "楷体": "KaiTi",
        "KaiTi": "KaiTi",
        "仿宋": "FangSong",
        "FangSong": "FangSong",
        "方正小标宋简体": "方正小标宋简体",  # Keep as is
        "方正小标宋简体": "方正小标宋简体",
    }

    def __init__(self, file_path: str):
        """
        Initialize the parser with a document file.

        Args:
            file_path: Path to the .docx file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        self.file_path = file_path
        self.doc = Document(file_path)

    def parse(self) -> Dict[str, Any]:
        """
        Parse the document and return structured data.

        Returns:
            Dictionary containing parsed document data
        """
        return {
            "info": self._parse_info(),
            "page_settings": self._parse_page_settings(),
            "paragraphs": self._parse_paragraphs(),
            "runs": self._parse_runs(),
            "headings": self._parse_headings(),
            "tables": self._parse_tables(),
            "figures": self._parse_figures(),
            "table_of_contents": self._parse_table_of_contents(),
            "heading_structure": self._parse_heading_structure(),
        }

    def _parse_info(self) -> Dict[str, Any]:
        """Parse basic document information."""
        file_size = os.path.getsize(self.file_path)
        filename = os.path.basename(self.file_path)

        return {
            "filename": filename,
            "file_size": file_size,
        }

    def _parse_page_settings(self) -> Dict[str, Any]:
        """Parse page settings including margins and paper size."""
        sections = self.doc.sections
        if not sections:
            return {}

        section = sections[0]

        # Convert EMU to mm
        def emu_to_mm(emu: Optional[float]) -> float:
            if emu is None:
                return 0.0
            return round((emu / 914400) * 25.4, 2)

        # Get page dimensions
        width_mm = emu_to_mm(section.page_width)
        height_mm = emu_to_mm(section.page_height)

        return {
            "paper_size": {
                "width_mm": width_mm,
                "height_mm": height_mm,
            },
            "margins": {
                "top_mm": emu_to_mm(section.top_margin),
                "bottom_mm": emu_to_mm(section.bottom_margin),
                "left_mm": emu_to_mm(section.left_margin),
                "right_mm": emu_to_mm(section.right_margin),
            },
        }

    def _parse_paragraphs(self) -> List[Dict[str, Any]]:
        """Parse all paragraphs in the document."""
        paragraphs = []

        # Estimate page info (assuming A4 with 1.5 line spacing, ~25 paragraphs per page)
        total_paragraphs = len(self.doc.paragraphs)
        paragraphs_per_page = 25  # Estimate

        # Track line numbers within current page
        current_page = 1
        line_in_page = 1
        para_start_line = {}  # Track starting line for each paragraph index

        for idx, para in enumerate(self.doc.paragraphs):
            # Estimate page number
            estimated_page = (idx // paragraphs_per_page) + 1

            # Count lines in this paragraph
            text = para.text.strip()
            line_count = max(1, len(text.split('\n')) + (len(text) // 50))  # Rough estimate

            # Store starting line
            para_start_line[idx] = line_in_page
            line_in_page += line_count

            # Check if we've moved to a new page
            if estimated_page > current_page:
                current_page = estimated_page
                line_in_page = 1

            # Get font info from the first run (for body text extraction)
            font_info = None
            if para.runs and para.runs[0].text:
                font_info = self._parse_font(para.runs[0].font)
            
            para_data = {
                "index": idx,
                "text": para.text,
                "style": {
                    "name": para.style.name if para.style else "Normal",
                },
                "formatting": self._parse_paragraph_formatting(para),
                "font": font_info or {},  # Add font info for template extraction
                "alignment": self._get_alignment(para.alignment),
                # Add page and line info
                "page_number": estimated_page,
                "start_line": para_start_line[idx],
                "end_line": line_in_page - 1,
                "line_count": line_count,
            }
            paragraphs.append(para_data)

        return paragraphs

    def _parse_paragraph_formatting(self, para) -> Dict[str, Any]:
        """Parse paragraph formatting."""
        formatting = {}

        if para.paragraph_format is not None:
            # First line indent
            first_line_indent = para.paragraph_format.first_line_indent
            if first_line_indent is not None:
                # Convert to mm, then estimate character count
                indent_mm = (first_line_indent.pt * 0.352778) if hasattr(first_line_indent, 'pt') else 0
                formatting["first_line_indent_mm"] = round(indent_mm, 2)
                # Estimate character count (assuming 12pt font, ~5.3mm per 2 chars)
                formatting["first_line_indent_chars"] = round(indent_mm / 2.65, 1)

            # Line spacing
            line_spacing = para.paragraph_format.line_spacing
            rule = para.paragraph_format.line_spacing_rule
            
            if line_spacing is not None:
                if isinstance(line_spacing, float):
                    # Multiple line spacing (e.g., 1.5 = 1.5x)
                    formatting["line_spacing"] = round(line_spacing, 2)
                    formatting["line_spacing_rule"] = "MULTIPLE"
                    # Convert to points: need font size, assume 12pt if not available
                    font_size = 12
                    if para.runs and para.runs[0].font.size:
                        font_size = para.runs[0].font.size.pt
                    formatting["line_spacing_pt"] = round(line_spacing * font_size, 2)
                else:
                    # Fixed line spacing in points
                    spacing_pt = round(line_spacing.pt, 2) if hasattr(line_spacing, 'pt') else None
                    if spacing_pt:
                        formatting["line_spacing"] = spacing_pt
                        formatting["line_spacing_pt"] = spacing_pt
                        formatting["line_spacing_rule"] = "EXACT"

            # Line spacing rule (if not already set)
            if rule is not None and "line_spacing_rule" not in formatting:
                formatting["line_spacing_rule"] = str(rule)

            # Space before/after
            space_before = para.paragraph_format.space_before
            if space_before is not None:
                formatting["space_before_pt"] = round(space_before.pt, 2) if hasattr(space_before, 'pt') else 0

            space_after = para.paragraph_format.space_after
            if space_after is not None:
                formatting["space_after_pt"] = round(space_after.pt, 2) if hasattr(space_after, 'pt') else 0

        return formatting

    def _parse_runs(self) -> List[Dict[str, Any]]:
        """Parse all text runs for font information."""
        runs = []
        paragraphs_per_page = 25

        for para_idx, para in enumerate(self.doc.paragraphs):
            # Estimate page and line info
            page_number = (para_idx // paragraphs_per_page) + 1
            text = para.text.strip() if para.text else ""
            line_count = max(1, len(text.split('\n')) + (len(text) // 50))

            for run_idx, run in enumerate(para.runs):
                if run.text:  # Only include runs with text
                    run_data = {
                        "paragraph_index": para_idx,
                        "run_index": run_idx,
                        "text": run.text,
                        "font": self._parse_font(run.font, para.style if para.style else None),
                        # Include page/line info
                        "page_number": page_number,
                    }
                    runs.append(run_data)

        return runs

    def _parse_font(self, font, style=None) -> Dict[str, Any]:
        """
        Parse font information.

        Args:
            font: The run's font object
            style: The paragraph's style object (for fallback when font properties are None)

        Returns:
            Dictionary with font properties
        """
        # Normalize font name: extract first font from comma-separated list
        # e.g., "SimSun,宋体" -> "SimSun"
        font_name = None
        size_pt = None
        bold = False

        # Try to get font name from run font
        if font.name:
            raw_name = font.name.split(',')[0].strip()
            font_name = self.FONT_NAME_MAP.get(raw_name, raw_name)

        # If no font name from run, try to get from style's font
        if not font_name and style and style.font:
            style_font = style.font
            if style_font.name:
                raw_name = style_font.name.split(',')[0].strip()
                font_name = self.FONT_NAME_MAP.get(raw_name, raw_name)

        # Try to get size from run font
        if font.size:
            size_pt = round(font.size.pt, 2)
        # If no size from run, try to get from style's font
        elif style and style.font and style.font.size:
            size_pt = round(style.font.size.pt, 2)

        # Try to get bold from run font
        if font.bold is not None:
            bold = font.bold
        # If no bold from run, try to get from style's font
        elif style and style.font and style.font.bold is not None:
            bold = style.font.bold

        font_data = {
            "name": font_name,
            "size_pt": size_pt,
            "bold": bold,
        }
        return font_data

    def _get_alignment(self, alignment) -> str:
        """Convert alignment enum to string."""
        if alignment is None:
            return "left"
        return self.ALIGNMENT_MAP.get(alignment, "left")

    def _parse_headings(self) -> List[Dict[str, Any]]:
        """Parse all headings in the document."""
        headings = []

        for para_idx, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            is_heading = False
            level = 1
            
            # Method 1: Check if style name starts with "Heading"
            if style_name.startswith("Heading"):
                is_heading = True
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
            
            # Method 2: Check visual characteristics (if no Heading style)
            # Titles often have: larger font size, bold, and are shorter lines
            elif para.runs:
                # Check first run for visual characteristics
                first_run = para.runs[0]
                font_size = first_run.font.size.pt if first_run.font.size else None
                is_bold = first_run.font.bold if first_run.font.bold is not None else False

                # Heuristic: If text is short, bold, and has larger font, it might be a heading
                # Use stricter criteria to avoid false positives
                # - Minimum font size of 14pt (up from 12pt)
                # - Maximum text length of 50 characters (down from 100)
                # - Must be bold
                if font_size and font_size >= 14 and is_bold:
                    # Check if text length suggests it's a heading (typically shorter than body text)
                    if len(text) <= 50:  # Headings are usually shorter
                        # Additional check: text should not be too long (like a full sentence)
                        # and should ideally start with a common heading pattern
                        is_heading = True
                        # Estimate level based on font size
                        if font_size >= 20:
                            level = 1
                        elif font_size >= 16:
                            level = 2
                        else:
                            level = 3
            
            if is_heading:
                heading_data = {
                    "level": level,
                    "text": text,
                    "style_name": style_name,
                    "alignment": self._get_alignment(para.alignment),
                    "paragraph_index": para_idx,
                }

                # Get font info from the first run
                if para.runs:
                    heading_data["font"] = self._parse_font(para.runs[0].font, para.style if para.style else None)

                headings.append(heading_data)

        return headings

    def _parse_tables(self) -> List[Dict[str, Any]]:
        """Parse all tables in the document and extract their captions."""
        tables = []

        # 获取所有段落索引对应的表格位置
        para_to_table = {}  # paragraph_index -> table_index
        for idx, table in enumerate(self.doc.tables):
            # 找到表格所在的段落索引
            # 表格通常嵌入在段落中，我们需要找到包含它的段落
            table_element = table._element
            parent = table_element.getparent()
            while parent is not None:
                if parent.tag.endswith('}p'):  # 段落元素
                    break
                parent = parent.getparent()

            # 找到表格在文档中的位置
            table_para_idx = None
            if parent is not None:
                # 尝试通过遍历段落找到表格位置
                for para_idx, para in enumerate(self.doc.paragraphs):
                    if para._element == parent:
                        table_para_idx = para_idx
                        break

            table_data = {
                "index": idx,
                "caption": None,
                "caption_paragraph_index": None,
                "caption_position": "above",  # 表格标题通常在上方
                "paragraph_index": table_para_idx,
            }

            # 查找表格标题（通常在表格前后的段落中）
            if table_para_idx is not None:
                caption = self._find_table_caption(table_para_idx, idx)
                if caption:
                    table_data["caption"] = caption["text"]
                    table_data["caption_paragraph_index"] = caption["para_index"]

            tables.append(table_data)

        return tables

    def _find_table_caption(self, table_para_idx: int, table_index: int) -> Optional[Dict[str, Any]]:
        """
        查找表格的标题。
        表格标题通常在表格前1-2个段落中，包含"表"字和编号。
        """
        # 检查表格前的段落（标题通常在上方）
        for offset in [-2, -1, 1, 2]:
            check_idx = table_para_idx + offset
            if 0 <= check_idx < len(self.doc.paragraphs):
                para = self.doc.paragraphs[check_idx]
                text = para.text.strip()

                # 匹配常见的表格标题格式：表1-1、表1.1、表1 等
                if re.match(r'^表\s*\d+([－\-.\．]\d+)*\s+', text):
                    return {
                        "text": text,
                        "para_index": check_idx
                    }

        return None

    def _parse_figures(self) -> List[Dict[str, Any]]:
        """Parse all figures/images in the document and extract their captions."""
        figures = []

        figure_index = 0
        figure_paras = []  # 记录包含图片的段落索引

        for para_idx, para in enumerate(self.doc.paragraphs):
            # Check if paragraph contains inline shapes (images)
            if hasattr(para, '_element'):
                # Use XPath to find blip elements (images)
                blips = para._element.xpath('.//a:blip')
                if blips:
                    # 查找图片标题
                    caption = self._find_figure_caption(para_idx, figure_index)

                    figure_data = {
                        "index": figure_index,
                        "caption": caption.get("text") if caption else None,
                        "caption_paragraph_index": caption.get("para_index") if caption else None,
                        "caption_position": "below",  # 图片标题通常在下方
                        "paragraph_index": para_idx,
                    }
                    figures.append(figure_data)
                    figure_paras.append(para_idx)
                    figure_index += 1

        return figures

    def _find_figure_caption(self, figure_para_idx: int, figure_index: int) -> Optional[Dict[str, Any]]:
        """
        查找图片的标题。
        图片标题通常在图片前后的段落中，包含"图"字和编号。
        """
        # 检查图片周围的段落
        for offset in [-2, -1, 1, 2]:
            check_idx = figure_para_idx + offset
            if 0 <= check_idx < len(self.doc.paragraphs):
                para = self.doc.paragraphs[check_idx]
                text = para.text.strip()

                # 匹配常见的图片标题格式：图1-1、图1.1、图1 等
                if re.match(r'^图\s*\d+([－\-.\．]\d+)*\s+', text):
                    return {
                        "text": text,
                        "para_index": check_idx
                    }

        return None

    def _parse_table_of_contents(self) -> Dict[str, Any]:
        """
        Parse table of contents (TOC) from the document.
        
        Returns:
            Dictionary with TOC information including entries
        """
        toc = {
            "exists": False,
            "entries": [],
            "paragraph_index": None,
        }
        
        # Pattern to match "目录" heading (may have spaces, numbers, etc.)
        # Support variations like "目录", "目  录", "一、目录", "1. 目录"
        toc_title_pattern = re.compile(r'^[\s　]*[一二三四五六七八九十\d\.、．]*[\s　]*目\s*录\s*$')
        
        # Find TOC title
        toc_start_idx = None
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            # Try exact match first
            if toc_title_pattern.match(text):
                toc["exists"] = True
                toc["paragraph_index"] = idx
                toc_start_idx = idx
                break
            # Also check if text contains "目录" and might be a TOC title
            # (sometimes TOC is just text without special formatting)
            elif "目录" in text and len(text) <= 10:  # Short text containing "目录"
                # Check if it's likely a title (not just body text mentioning "目录")
                normalized = re.sub(r'[\s　\d\.、．一二三四五六七八九十]+', '', text)
                if normalized == "目录":
                    toc["exists"] = True
                    toc["paragraph_index"] = idx
                    toc_start_idx = idx
                    break
        
        if toc_start_idx is not None:
            # Extract TOC entries (usually found after the TOC title)
            toc["entries"] = self._extract_toc_entries(toc_start_idx)
        
        return toc

    def _extract_toc_entries(self, toc_start_idx: int) -> List[Dict[str, Any]]:
        """
        Extract table of contents entries.
        
        TOC entries typically have:
        - Title text
        - Tab characters or dots
        - Page numbers
        
        Args:
            toc_start_idx: Index of the paragraph containing "目录"
            
        Returns:
            List of TOC entry dictionaries
        """
        entries = []
        
        # Look for TOC entries in the next 100 paragraphs (reasonable limit)
        # TOC entries typically end when we find a heading or substantial content
        max_scan = min(toc_start_idx + 100, len(self.doc.paragraphs))
        
        # Patterns to identify TOC entries:
        # 1. Contains tab characters (\t)
        # 2. Ends with a page number
        # 3. May have dots or underscores as separators
        
        page_num_pattern = re.compile(r'(\d+)\s*$')  # Page number at end
        
        for idx in range(toc_start_idx + 1, max_scan):
            para = self.doc.paragraphs[idx]
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Skip if we hit a major heading (likely end of TOC)
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading") and para.style.name in ["Heading 1", "1. 标题 1"]:
                # Check if this looks like content, not TOC
                if len(text) > 10:  # Substantial heading text
                    break
            
            # Check if this looks like a TOC entry
            # TOC entries often contain:
            # - Tab characters
            # - Dots/underscores before page number
            # - Page number at the end
            
            if '\t' in text or re.search(r'\.{2,}', text) or re.search(r'_{2,}', text):
                # Extract title and page number
                # Remove page number and separators
                clean_text = re.sub(r'[\s\t\._]+[\d]+$', '', text)  # Remove page number
                clean_text = clean_text.strip().rstrip('.\t_')
                
                # Try to extract page number
                page_match = page_num_pattern.search(text)
                page_number = int(page_match.group(1)) if page_match else None
                
                # Determine level by indentation or style
                # Level 1: no indentation, Level 2: some indentation, etc.
                level = 1
                if text.startswith('\t'):
                    level = 2
                elif text.startswith('\t\t'):
                    level = 3
                
                # Check paragraph style for heading level
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        pass
                
                if clean_text:  # Only add non-empty entries
                    entries.append({
                        "title": clean_text,
                        "level": level,
                        "page_number": page_number,
                        "paragraph_index": idx,
                        "original_text": text,
                    })
        
        return entries

    def _parse_heading_structure(self) -> Dict[str, Any]:
        """
        Parse complete heading hierarchy structure.
        
        Returns:
            Dictionary with heading tree and flat list
        """
        headings = self._parse_headings()
        
        # Build tree structure
        tree = []
        flat = []
        stack = []  # Stack to track parent headings at each level
        
        for heading in headings:
            level = heading.get("level", 1)
            text = heading.get("text", "").strip()
            
            # Build flat list with hierarchy info
            flat_entry = {
                "level": level,
                "text": text,
                "paragraph_index": heading.get("paragraph_index"),
                "parent": None,
                "children": []
            }
            
            # Pop stack until we find parent at lower level
            while stack and stack[-1]["level"] >= level:
                stack.pop()
            
            # Set parent if stack is not empty
            if stack:
                flat_entry["parent"] = stack[-1]["text"]
            
            # Add to parent's children if exists
            if stack:
                parent_idx = None
                for i, item in enumerate(flat):
                    if item["text"] == stack[-1]["text"]:
                        parent_idx = i
                        break
                if parent_idx is not None:
                    flat[parent_idx]["children"].append(text)
            
            flat.append(flat_entry)
            stack.append({"level": level, "text": text})
        
        # Build tree structure (nested)
        if flat:
            tree = self._build_heading_tree(flat)
        
        return {
            "tree": tree,
            "flat": flat,
        }

    def _build_heading_tree(self, flat_headings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Build nested tree structure from flat heading list.
        
        Args:
            flat_headings: Flat list of headings with parent info
            
        Returns:
            Nested tree structure
        """
        if not flat_headings:
            return []
        
        tree = []
        stack = []  # Stack of (node, level) tuples
        
        for heading in flat_headings:
            level = heading["level"]
            node = {
                "level": level,
                "text": heading["text"],
                "paragraph_index": heading["paragraph_index"],
                "children": []
            }
            
            # Pop stack until we find parent
            while stack and stack[-1][1] >= level:
                stack.pop()
            
            # Add to parent or root
            if stack:
                stack[-1][0]["children"].append(node)
            else:
                tree.append(node)
            
            stack.append((node, level))
        
        return tree


def parse_document_safe(file_path: str) -> Dict[str, Any]:
    """
    Safely parse a document with error handling.

    Args:
        file_path: Path to the document file

    Returns:
        Dictionary with success status and data or error
    """
    try:
        parser = DocxParser(file_path)
        data = parser.parse()
        return {"success": True, "data": data}
    except FileNotFoundError as e:
        return {"success": False, "error": str(e), "error_type": "FileNotFoundError"}
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }


# ============== Utility Functions ==============

def contains_chinese(text: str) -> bool:
    """Check if text contains Chinese characters."""
    return bool(re.search(r'[\u4e00-\u9fff]', text))


def contains_english(text: str) -> bool:
    """Check if text contains English characters."""
    return bool(re.search(r'[a-zA-Z]', text))
