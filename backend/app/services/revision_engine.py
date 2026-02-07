"""
Revision Engine
Generates a revised version of the document based on check results.
"""
import os
import shutil
import uuid
import logging
from typing import List, Dict, Any
from docx import Document
from app.utils.xml_reviser import XmlReviser

logger = logging.getLogger(__name__)

class RevisionEngine:
    """Engine to apply fixes to a document."""
    
    def __init__(self, original_file_path: str, output_dir: str):
        self.original_file_path = original_file_path
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
    def generate_revised_document(self, issues: List[Dict[str, Any]]) -> str:
        """
        Generate a revised document with track changes enabled.
        
        Args:
            issues: List of issues detected by RuleEngine
            
        Returns:
            Path to the generated revised file
        """
        # 1. Create a copy of the original file
        filename = os.path.basename(self.original_file_path)
        name, ext = os.path.splitext(filename)
        revised_filename = f"{name}_revised_{uuid.uuid4().hex[:8]}{ext}"
        revised_path = os.path.join(self.output_dir, revised_filename)
        shutil.copy2(self.original_file_path, revised_path)
        
        # 2. Load the copy
        doc = Document(revised_path)
        reviser = XmlReviser(doc)
        
        # 3. Enable track changes
        reviser.enable_track_revisions()
        
        # 4. Iterate all issues
        paragraphs = doc.paragraphs
        
        # Debug: Log issues count and structure
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"Processing {len(issues)} issues for revision")
        
        issues_with_fix = [i for i in issues if i.get("fix_action")]
        logger.info(f"Found {len(issues_with_fix)} issues with fix_action")
        
        for issue in issues:
            # Use raw_locations if available (from RuleEngine grouping), otherwise fallback to single location
            # RuleEngine now returns 'raw_locations' list for grouped issues
            locations = issue.get("raw_locations")
            if not locations:
                logger.debug(f"Issue {issue.get('rule_id')} has no raw_locations, using fallback")
                # Fallback to single location if not grouped or old format
                single_location = issue.get("location", {})
                # If location type is "merged", we can't use it directly
                # In this case, we should skip or try to infer from issue category
                if single_location.get("type") == "merged":
                    # For merged locations, check if it's a document-level fix
                    # or if we can infer from the issue category
                    if issue.get("category") == "page":
                        # Page-level issues are document-level fixes
                        self._apply_fix(reviser, issue, None)
                        continue
                    else:
                        # For other merged locations, we can't process them
                        # Log warning and skip
                        logger.warning(f"Skipping issue {issue.get('rule_id')} with merged location and no raw_locations")
                        continue
                locations = [single_location]
            
            # Apply fix for EACH location
            applied_count = 0
            for loc in locations:
                if not loc: continue
                
                # Determine target object
                target_obj = None
                
                # Check location type
                loc_type = loc.get("type")
                
                if loc_type == "document":
                    # Document level fix
                    # For document level, we only need to apply once per issue group ideally, 
                    # but applying idempotent fixes multiple times is safe.
                    logger.debug(f"Applying document-level fix for {issue.get('rule_id')}")
                    self._apply_fix(reviser, issue, None)
                    applied_count += 1
                    continue 
                
                elif loc_type in ["paragraph", "heading"]:
                    # Find paragraph
                    para_index = None
                    if "index" in loc:
                        para_index = loc["index"]
                    elif "paragraph_index" in loc:
                        para_index = loc["paragraph_index"]
                    
                    if para_index is not None and 0 <= para_index < len(paragraphs):
                        target_obj = paragraphs[para_index]
                        logger.debug(f"Found target paragraph {para_index} for {issue.get('rule_id')}")
                    else:
                        logger.warning(f"Could not find paragraph {para_index} for {issue.get('rule_id')} (total paragraphs: {len(paragraphs)})")
                
                # Apply fix to target (if found, or if it's a doc-level fix handled above)
                if target_obj:
                    logger.debug(f"Applying fix {issue.get('fix_action')} to paragraph {para_index}")
                    self._apply_fix(reviser, issue, target_obj)
                    applied_count += 1
            
            if applied_count == 0:
                logger.warning(f"No fixes applied for issue {issue.get('rule_id')} (fix_action: {issue.get('fix_action')})")

        # 5. Save
        doc.save(revised_path)
        return revised_path

    def _apply_fix(self, reviser: XmlReviser, issue: Dict[str, Any], target_obj: Any):
        """Apply a single fix."""
        action = issue.get("fix_action")
        params = issue.get("fix_params") or {}
        suggestion = issue.get("suggestion", "")

        # If no automatic fix action, just add a comment/annotation
        if not action:
            if target_obj and hasattr(target_obj, 'add_run'): # Is a paragraph
                reviser.add_comment(target_obj, f"建议: {suggestion}")
            return

        try:
            # ========== 格式类修复 ==========
            if action == "set_page_margin":
                reviser.set_page_margin(
                    params.get("top_mm", 25.4),
                    params.get("bottom_mm", 25.4),
                    params.get("left_mm", 31.8),
                    params.get("right_mm", 31.8)
                )

            elif action == "set_paragraph_indent" and target_obj:
                reviser.set_paragraph_indent(target_obj, params.get("first_line_indent_chars", 2))

            elif action == "set_heading_style" and target_obj:
                # Need to determine level to pick right params
                # Issue location might have level, or params might have levels dict
                # HEADING_STYLES rule params structure: { "level1": {...}, "level2": {...} }
                # We need to know which level this paragraph is.
                # Assuming target_obj is a paragraph, we check its style
                style_name = target_obj.style.name
                level = 1 # default
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                    except:
                        pass

                level_key = f"level{level}"
                if level_key in params:
                    style_params = params[level_key]
                    reviser.set_heading_style(
                        target_obj,
                        style_params.get("font", "SimHei"),
                        style_params.get("size_pt", 14),
                        style_params.get("alignment", "left"),
                        style_params.get("bold", True)
                    )

            elif action == "set_run_style" and target_obj:
                # Reuse set_heading_style logic but for body text (no alignment change usually, or explicit)
                # For body text, we usually want to set font name and size.
                # Note: set_heading_style applies to all runs in paragraph.
                reviser.set_heading_style(
                    target_obj,
                    params.get("chinese_font", "宋体"), # Default fallback
                    params.get("size_pt", 12),
                    None, # Do not change alignment
                    False # Not bold
                )

            elif action == "set_page_size":
                 # Not implemented in reviser yet, but harmless to skip or add comment
                 pass

            # ========== AI内容类修复 ==========
            elif action == "replace_text" and target_obj:
                # 替换错别字
                original = params.get("original", "")
                correction = params.get("correction", "")

                if original and correction and hasattr(target_obj, 'runs'):
                    # Try to replace in the paragraph
                    success = reviser.replace_text_in_paragraph(target_obj, original, correction)
                    if not success:
                        # If paragraph-level replacement failed, try run-level
                        for run in target_obj.runs:
                            if original in run.text:
                                reviser.replace_text_in_run(run, original, correction)
                                break

            elif action == "replace_ref" and target_obj:
                # 替换引用编号
                original_ref = params.get("original_ref", "")
                suggested_ref = params.get("suggested_ref", "")

                if original_ref and suggested_ref and hasattr(target_obj, 'runs'):
                    for run in target_obj.runs:
                        if original_ref in run.text:
                            reviser.replace_text_in_run(run, original_ref, suggested_ref)
                            break

            # Always add a comment explaining what was done
            if target_obj and hasattr(target_obj, 'runs'):
                 # Note: add_comment is currently a no-op, but kept for future use
                 reviser.add_comment(target_obj, f"已自动修正: {suggestion}")

        except Exception as e:
            print(f"Failed to apply fix {action}: {e}")
            # Fallback to comment
            if target_obj:
                reviser.add_comment(target_obj, f"自动修正失败: {suggestion}")

