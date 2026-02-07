"""
Pytest configuration and fixtures for backend tests.
"""
import pytest
import os
import tempfile
import shutil
from datetime import date, datetime
from typing import Generator
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
from fastapi.testclient import TestClient

from app.core.database import Base, get_db
from app.core.security import get_password_hash
from app.models import User, Check, CheckType, CheckStatus, CostType, RuleTemplate
from main import app


# Test database setup
TEST_DATABASE_URL = "sqlite:///./test.db"
engine = create_engine(TEST_DATABASE_URL, connect_args={"check_same_thread": False})
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


@pytest.fixture(scope="session")
def test_db_engine():
    """Create test database engine."""
    Base.metadata.create_all(bind=engine)
    yield engine
    Base.metadata.drop_all(bind=engine)


@pytest.fixture(scope="function")
def db(test_db_engine) -> Generator[Session, None, None]:
    """Create a new database session for each test."""
    connection = test_db_engine.connect()
    transaction = connection.begin()
    session = TestingSessionLocal(bind=connection)

    yield session

    session.close()
    transaction.rollback()
    connection.close()


@pytest.fixture(scope="function")
def client(db: Session) -> Generator[TestClient, None, None]:
    """Create a test client with database session override."""
    def override_get_db():
        try:
            yield db
        finally:
            pass

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


@pytest.fixture
def test_user(db: Session) -> User:
    """Create a test user."""
    user = User(
        username="testuser",
        password_hash=get_password_hash("testpass123"),
        nickname="Test User",
        free_count=3,
        basic_count=10,
        full_count=5,
        last_reset_date=date.today()
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@pytest.fixture
def guest_user(db: Session) -> User:
    """Create a guest user."""
    user = User(
        username="guest_test123",
        password_hash=get_password_hash("guest_password"),
        nickname="Guest User",
        free_count=1,
        basic_count=0,
        full_count=0,
        last_reset_date=date.today()
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@pytest.fixture
def auth_token(client: TestClient, test_user: User) -> str:
    """Get authentication token for test user."""
    response = client.post(
        "/api/auth/login",
        json={
            "username": "testuser",
            "password": "testpass123"
        }
    )
    assert response.status_code == 200
    data = response.json()
    return data["data"]["access_token"]


@pytest.fixture
def auth_headers(auth_token: str) -> dict:
    """Get authentication headers with bearer token."""
    return {"Authorization": f"Bearer {auth_token}"}


@pytest.fixture
def temp_upload_dir() -> Generator[str, None, None]:
    """Create a temporary upload directory."""
    temp_dir = tempfile.mkdtemp()
    yield temp_dir
    shutil.rmtree(temp_dir, ignore_errors=True)


@pytest.fixture
def sample_docx_path(temp_upload_dir: str) -> str:
    """Create a sample .docx file for testing."""
    from docx import Document

    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is section 1 content.')

    file_path = os.path.join(temp_upload_dir, "test_document.docx")
    doc.save(file_path)
    return file_path


@pytest.fixture
def sample_doc_data() -> dict:
    """Create sample parsed document data."""
    return {
        "info": {
            "filename": "test.docx",
            "file_size": 12345
        },
        "page_settings": {
            "paper_size": {
                "width_mm": 210.0,
                "height_mm": 297.0
            },
            "margins": {
                "top_mm": 25.4,
                "bottom_mm": 25.4,
                "left_mm": 31.7,
                "right_mm": 31.7
            }
        },
        "paragraphs": [
            {
                "index": 0,
                "text": "Test Title",
                "style": {"name": "Heading 1"},
                "formatting": {},
                "font": {"name": "SimHei", "size_pt": 18, "bold": True},
                "alignment": "center",
                "page_number": 1,
                "start_line": 1,
                "end_line": 1,
                "line_count": 1
            },
            {
                "index": 1,
                "text": "This is body text.",
                "style": {"name": "Normal"},
                "formatting": {
                    "first_line_indent_chars": 2.0,
                    "line_spacing_pt": 28.0
                },
                "font": {"name": "SimSun", "size_pt": 14, "bold": False},
                "alignment": "left",
                "page_number": 1,
                "start_line": 2,
                "end_line": 2,
                "line_count": 1
            }
        ],
        "runs": [
            {
                "paragraph_index": 0,
                "run_index": 0,
                "text": "Test Title",
                "font": {"name": "SimHei", "size_pt": 18, "bold": True},
                "page_number": 1
            },
            {
                "paragraph_index": 1,
                "run_index": 0,
                "text": "This is body text.",
                "font": {"name": "SimSun", "size_pt": 14, "bold": False},
                "page_number": 1
            }
        ],
        "headings": [
            {
                "level": 1,
                "text": "Test Title",
                "style_name": "Heading 1",
                "alignment": "center",
                "paragraph_index": 0,
                "font": {"name": "SimHei", "size_pt": 18, "bold": True}
            }
        ],
        "tables": [],
        "figures": [],
        "table_of_contents": {
            "exists": False,
            "entries": [],
            "paragraph_index": None
        },
        "heading_structure": {
            "tree": [],
            "flat": []
        }
    }


@pytest.fixture
def sample_rules() -> list:
    """Create sample rules for testing."""
    return [
        {
            "id": "FONT_CHECK_BODY",
            "name": "正文字体检查",
            "category": "font",
            "match": "run",
            "condition": {
                "chinese_font": "SimSun",
                "chinese_size_pt": 14
            },
            "error_message": "正文应使用宋体14pt",
            "suggestion": "请将正文字体设置为宋体，字号14pt",
            "checker": "deterministic"
        },
        {
            "id": "MARGIN_CHECK",
            "name": "页边距检查",
            "category": "page",
            "match": "document",
            "condition": {
                "top_mm": 25.4,
                "bottom_mm": 25.4,
                "left_mm": 31.7,
                "right_mm": 31.7,
                "tolerance_mm": 2
            },
            "error_message": "页边距不符合规范",
            "suggestion": "请设置页边距为：上25.4mm，下25.4mm，左31.7mm，右31.7mm",
            "checker": "deterministic"
        }
    ]


@pytest.fixture
def sample_rule_template(db: Session, test_user: User) -> RuleTemplate:
    """Create a sample rule template."""
    from app.models.rule_template import TemplateType
    template = RuleTemplate(
        name="测试模板",
        description="用于测试的规则模板",
        template_type=TemplateType.SYSTEM,
        user_id=None,  # System template has no user
        is_default=True,
        config_json={
            "page": {
                "margins": {
                    "top_cm": 2.5,
                    "bottom_cm": 2.5,
                    "left_cm": 3.0,
                    "right_cm": 2.5
                },
                "paper_name": "A4"
            },
            "body": {
                "font": "SimSun",
                "size_pt": 14,
                "line_spacing_pt": 28,
                "first_line_indent_chars": 2
            },
            "headings": [
                {
                    "level": 1,
                    "font": "SimHei",
                    "size_pt": 18,
                    "bold": True,
                    "alignment": "center"
                }
            ]
        }
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    return template


@pytest.fixture
def sample_check(db: Session, test_user: User, sample_docx_path: str) -> Check:
    """Create a sample check record."""
    check = Check(
        check_id="check_test123",
        user_id=test_user.id,
        file_id="file_test123",
        filename="test_document.docx",
        file_path=sample_docx_path,
        check_type=CheckType.BASIC,
        status=CheckStatus.COMPLETED,
        cost_type=CostType.FREE,
        result_json='{"total_issues": 2, "issues": []}'
    )
    db.add(check)
    db.commit()
    db.refresh(check)
    return check
