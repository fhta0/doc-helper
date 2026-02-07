"""
Integration tests for Rule Template API.
"""
import pytest
import os
import tempfile
from io import BytesIO
from sqlalchemy.orm import Session
from app.models import RuleTemplate, TemplateType
from docx import Document


@pytest.mark.integration
@pytest.mark.api
class TestRuleTemplatesAPI:
    """Test cases for rule template endpoints."""

    def test_get_rule_templates_all(self, client, db: Session, test_user, sample_rule_template):
        """Test getting all rule templates (system + user's custom)."""
        response = client.get("/api/rule-templates")

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        assert len(data["data"]["templates"]) > 0

    def test_get_rule_templates_system_only(self, client):
        """Test getting only system templates."""
        response = client.get("/api/rule-templates?template_type=system")

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        # All templates should be system type
        for template in data["data"]["templates"]:
            assert template["template_type"] == "system"

    def test_get_rule_templates_custom_authenticated(self, client, auth_headers, test_user, db: Session):
        """Test getting custom templates when authenticated."""
        # Create a custom template
        custom_template = RuleTemplate(
            name="Custom Template",
            description="User's custom template",
            template_type=TemplateType.CUSTOM,
            user_id=test_user.id,
            config_json={"page": {"margins": {"top_cm": 2.5}}}
        )
        db.add(custom_template)
        db.commit()

        response = client.get(
            "/api/rule-templates?template_type=custom",
            headers=auth_headers
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        # Should have at least the custom template
        assert any(t["name"] == "Custom Template" for t in data["data"]["templates"])

    def test_get_rule_template_by_id(self, client, sample_rule_template):
        """Test getting a specific template by ID."""
        response = client.get(f"/api/rule-templates/{sample_rule_template.id}")

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        assert data["data"]["id"] == sample_rule_template.id
        assert data["data"]["name"] == "测试模板"

    def test_get_rule_template_not_found(self, client):
        """Test getting a non-existent template."""
        response = client.get("/api/rule-templates/99999")

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 1004
        assert "不存在" in data["message"]

    def test_create_rule_template(self, client, auth_headers, test_user, db: Session):
        """Test creating a custom rule template."""
        request_data = {
            "name": "My Custom Template",
            "description": "A template for testing",
            "config": {
                "page": {
                    "margins": {
                        "top_cm": 2.5,
                        "bottom_cm": 2.5,
                        "left_cm": 3.0,
                        "right_cm": 2.5
                    }
                },
                "body": {
                    "font": "SimSun",
                    "size_pt": 14
                }
            }
        }

        response = client.post(
            "/api/rule-templates",
            headers=auth_headers,
            json=request_data
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        assert data["data"]["name"] == "My Custom Template"
        assert data["data"]["template_type"] == "custom"

        # Verify in database
        template = db.query(RuleTemplate).filter(RuleTemplate.name == "My Custom Template").first()
        assert template is not None
        assert template.user_id == test_user.id

    def test_create_rule_template_unauthenticated(self, client):
        """Test that unauthenticated users cannot create templates."""
        request_data = {
            "name": "Should Fail",
            "description": "Test",
            "config": {}
        }

        response = client.post("/api/rule-templates", json=request_data)

        assert response.status_code in [401, 403]

    def test_update_rule_template_custom(self, client, auth_headers, test_user, db: Session):
        """Test updating a custom template."""
        # Create a custom template
        template = RuleTemplate(
            name="Original Name",
            description="Original description",
            template_type=TemplateType.CUSTOM,
            user_id=test_user.id,
            config_json={}
        )
        db.add(template)
        db.commit()

        update_data = {
            "name": "Updated Name",
            "description": "Updated description"
        }

        response = client.put(
            f"/api/rule-templates/{template.id}",
            headers=auth_headers,
            json=update_data
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        assert data["data"]["name"] == "Updated Name"

        # Verify in database
        db.refresh(template)
        assert template.name == "Updated Name"
        assert template.description == "Updated description"

    def test_update_rule_template_system_forbidden(self, client, auth_headers, sample_rule_template):
        """Test that system templates cannot be updated."""
        update_data = {
            "name": "Hacked Name"
        }

        response = client.put(
            f"/api/rule-templates/{sample_rule_template.id}",
            headers=auth_headers,
            json=update_data
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 1003
        assert "不能修改" in data["message"]

    def test_update_rule_template_other_user_forbidden(self, client, auth_headers, db: Session, test_user):
        """Test that users cannot update other users' templates."""
        from app.models import User
        from app.core.security import get_password_hash

        # Create another user
        other_user = User(
            username="other_user",
            password_hash=get_password_hash("password"),
            nickname="Other"
        )
        db.add(other_user)
        db.commit()

        # Create template for other user
        template = RuleTemplate(
            name="Other's Template",
            template_type=TemplateType.CUSTOM,
            user_id=other_user.id,
            config_json={}
        )
        db.add(template)
        db.commit()

        update_data = {"name": "Stolen"}

        response = client.put(
            f"/api/rule-templates/{template.id}",
            headers=auth_headers,
            json=update_data
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 1003
        assert "无权" in data["message"]

    def test_delete_rule_template_custom(self, client, auth_headers, test_user, db: Session):
        """Test deleting a custom template."""
        template = RuleTemplate(
            name="To Delete",
            template_type=TemplateType.CUSTOM,
            user_id=test_user.id,
            config_json={}
        )
        db.add(template)
        db.commit()
        template_id = template.id

        response = client.delete(
            f"/api/rule-templates/{template_id}",
            headers=auth_headers
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200

        # Verify deleted
        deleted_template = db.query(RuleTemplate).filter(RuleTemplate.id == template_id).first()
        assert deleted_template is None

    def test_delete_rule_template_system_forbidden(self, client, auth_headers, sample_rule_template):
        """Test that system templates cannot be deleted."""
        response = client.delete(
            f"/api/rule-templates/{sample_rule_template.id}",
            headers=auth_headers
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 1003
        assert "不能删除" in data["message"]

    def test_delete_rule_template_other_user_forbidden(self, client, auth_headers, db: Session, test_user):
        """Test that users cannot delete other users' templates."""
        from app.models import User
        from app.core.security import get_password_hash

        # Create another user
        other_user = User(
            username="another_user",
            password_hash=get_password_hash("password"),
            nickname="Another"
        )
        db.add(other_user)
        db.commit()

        # Create template for other user
        template = RuleTemplate(
            name="Protected Template",
            template_type=TemplateType.CUSTOM,
            user_id=other_user.id,
            config_json={}
        )
        db.add(template)
        db.commit()

        response = client.delete(
            f"/api/rule-templates/{template.id}",
            headers=auth_headers
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 1003
        assert "无权" in data["message"]

    def test_parse_docx_to_rule(self, client, auth_headers, temp_upload_dir):
        """Test parsing docx file to extract rule configuration."""
        # Create a sample docx with specific formatting
        doc = Document()

        # Set page margins
        section = doc.sections[0]
        from docx.shared import Mm
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(32)
        section.right_margin = Mm(25)

        # Add headings with different styles
        heading1 = doc.add_heading('第一章 绪论', 1)
        heading1.alignment = 0  # Center (not working in test, but structure is there)

        doc.add_heading('第一节 研究背景', 2)

        # Add body paragraph
        para = doc.add_paragraph('这是正文内容，用于测试格式提取功能。')
        from docx.shared import Pt
        para.runs[0].font.name = '宋体'
        para.runs[0].font.size = Pt(14)

        file_path = os.path.join(temp_upload_dir, "sample_format.docx")
        doc.save(file_path)

        # Upload file for parsing
        with open(file_path, "rb") as f:
            response = client.post(
                "/api/rule-templates/parse/docx",
                headers=auth_headers,
                files={"file": ("sample_format.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
        assert "config" in data["data"]
        assert "structure" in data["data"]

        # Check config structure
        config = data["data"]["config"]
        assert "page" in config
        assert "body" in config

    def test_parse_text_to_rule(self, client, auth_headers):
        """Test parsing natural language text to rule configuration."""
        text = "要求使用宋体14磅字，行距28磅，页边距上下2.5厘米，左右3厘米"

        response = client.post(
            f"/api/rule-templates/parse/text?text={text}",
            headers=auth_headers
        )

        # This might fail if AI is not configured
        # But should return a valid response structure
        assert response.status_code == 200
        data = response.json()
        # Either success with config or error message
        assert "code" in data

    def test_template_config_json_persistence(self, client, auth_headers, test_user, db: Session):
        """Test that config JSON is properly stored and retrieved."""
        complex_config = {
            "page": {
                "margins": {
                    "top_cm": 2.54,
                    "bottom_cm": 2.54,
                    "left_cm": 3.17,
                    "right_cm": 3.17
                },
                "paper_name": "A4"
            },
            "body": {
                "font": "SimSun",
                "size_pt": 14,
                "line_spacing_pt": 28,
                "first_line_indent_chars": 2
            },
            "headings": [
                {
                    "level": 1,
                    "font": "SimHei",
                    "size_pt": 18,
                    "bold": True,
                    "alignment": "center"
                },
                {
                    "level": 2,
                    "font": "KaiTi",
                    "size_pt": 16,
                    "bold": False,
                    "alignment": "left"
                }
            ]
        }

        # Create template
        create_response = client.post(
            "/api/rule-templates",
            headers=auth_headers,
            json={
                "name": "Complex Template",
                "description": "Test complex config",
                "config": complex_config
            }
        )

        template_id = create_response.json()["data"]["id"]

        # Retrieve template
        get_response = client.get(f"/api/rule-templates/{template_id}")

        # Note: to_dict returns "config" not "config_json"
        assert get_response.json()["data"]["config"] == complex_config

    def test_templates_without_auth(self, client):
        """Test getting system templates without authentication."""
        response = client.get("/api/rule-templates")

        # Should allow access to system templates
        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 200
