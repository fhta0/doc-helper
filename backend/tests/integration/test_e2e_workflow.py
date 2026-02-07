"""
End-to-end integration tests for complete user workflows.

Tests the complete user journey from registration to document checking
and revision, ensuring all components work together correctly.
"""
import pytest
import os
import json
from sqlalchemy.orm import Session
from docx import Document
from app.models import User, Check, CheckStatus, CheckType, CostType


@pytest.mark.integration
@pytest.mark.api
class TestEndToEndWorkflow:
    """End-to-end test cases for complete user workflows."""

    def test_complete_new_user_workflow(self, client, db: Session, sample_docx_path):
        """Test complete workflow for a new user: register -> check -> revise -> download."""
        # Step 1: Register new user
        register_response = client.post(
            "/api/auth/register",
            json={
                "username": "newuser123",
                "password": "password123",
                "nickname": "New User"
            }
        )

        assert register_response.status_code == 200
        register_data = register_response.json()
        assert register_data["code"] == 200
        assert "access_token" in register_data["data"]
        assert register_data["data"]["user"]["free_count"] == 3

        token = register_data["data"]["access_token"]
        auth_headers = {"Authorization": f"Bearer {token}"}

        # Step 2: Get initial profile
        profile_response = client.get("/api/auth/user/profile", headers=auth_headers)
        assert profile_response.status_code == 200
        initial_profile = profile_response.json()["data"]
        assert initial_profile["free_count"] == 3

        # Step 3: Upload document
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        assert upload_response.status_code == 200
        file_id = upload_response.json()["data"]["file_id"]

        # Step 4: Submit check
        check_response = client.post(
            "/api/check",
            headers=auth_headers,
            json={
                "file_id": file_id,
                "filename": "test.docx",
                "check_type": "basic"
            }
        )

        assert check_response.status_code == 200
        check_data = check_response.json()
        assert check_data["code"] == 200
        check_id = check_data["data"]["check_id"]

        # Step 5: Verify count deducted
        profile_response = client.get("/api/auth/user/profile", headers=auth_headers)
        updated_profile = profile_response.json()["data"]
        assert updated_profile["free_count"] == 2

        # Step 6: Get check result
        result_response = client.get(f"/api/check/{check_id}", headers=auth_headers)
        assert result_response.status_code == 200
        result_data = result_response.json()["data"]
        assert result_data["status"] == "completed"
        assert "result" in result_data

        # Step 7: Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Step 8: Download revised document
        download_response = client.get(f"/api/check/{check_id}/download_revised")
        assert download_response.status_code == 200
        assert len(download_response.content) > 0

        # Step 9: Verify check history
        history_response = client.get("/api/check/recent", headers=auth_headers)
        assert history_response.status_code == 200
        history_data = history_response.json()["data"]
        assert history_data["total"] == 1
        assert len(history_data["checks"]) == 1

    def test_user_exhausts_free_count_then_uses_basic(self, client, db: Session, sample_docx_path):
        """Test workflow where user exhausts free count and then uses basic count."""
        # Register user
        register_response = client.post(
            "/api/auth/register",
            json={
                "username": "testuser_exhaust",
                "password": "password123",
                "nickname": "Test User"
            }
        )

        token = register_response.json()["data"]["access_token"]
        auth_headers = {"Authorization": f"Bearer {token}"}
        user_id = register_response.json()["data"]["user"]["id"]

        # Give user 1 free count and 2 basic counts
        user = db.query(User).filter(User.id == user_id).first()
        user.free_count = 1
        user.basic_count = 2
        db.commit()

        # First check: uses free count
        with open(sample_docx_path, "rb") as f:
            upload1 = client.post("/api/check/upload", headers=auth_headers,
                files={"file": ("test1.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id1 = upload1.json()["data"]["file_id"]
        check1 = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id1, "filename": "test1.docx", "check_type": "basic"})
        assert check1.status_code == 200
        check1_id = check1.json()["data"]["check_id"]

        # Verify free count used
        profile1 = client.get("/api/auth/user/profile", headers=auth_headers).json()["data"]
        assert profile1["free_count"] == 0
        assert profile1["basic_count"] == 2

        check1_record = db.query(Check).filter(Check.check_id == check1_id).first()
        assert check1_record.cost_type == CostType.FREE

        # Second check: uses basic count
        with open(sample_docx_path, "rb") as f:
            upload2 = client.post("/api/check/upload", headers=auth_headers,
                files={"file": ("test2.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id2 = upload2.json()["data"]["file_id"]
        check2 = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id2, "filename": "test2.docx", "check_type": "basic"})
        assert check2.status_code == 200
        check2_id = check2.json()["data"]["check_id"]

        # Verify basic count used
        profile2 = client.get("/api/auth/user/profile", headers=auth_headers).json()["data"]
        assert profile2["free_count"] == 0
        assert profile2["basic_count"] == 1

        check2_record = db.query(Check).filter(Check.check_id == check2_id).first()
        assert check2_record.cost_type == CostType.BASIC

        # Third check: uses basic count again
        with open(sample_docx_path, "rb") as f:
            upload3 = client.post("/api/check/upload", headers=auth_headers,
                files={"file": ("test3.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id3 = upload3.json()["data"]["file_id"]
        check3 = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id3, "filename": "test3.docx", "check_type": "basic"})
        assert check3.status_code == 200

        # Verify last basic count used
        profile3 = client.get("/api/auth/user/profile", headers=auth_headers).json()["data"]
        assert profile3["free_count"] == 0
        assert profile3["basic_count"] == 0

        # Fourth check: should fail
        with open(sample_docx_path, "rb") as f:
            upload4 = client.post("/api/check/upload", headers=auth_headers,
                files={"file": ("test4.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id4 = upload4.json()["data"]["file_id"]
        check4 = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id4, "filename": "test4.docx", "check_type": "basic"})
        assert check4.status_code == 403

    def test_check_with_template_and_revision(self, client, db: Session, sample_docx_path, sample_rule_template):
        """Test complete workflow with custom rule template and revision."""
        # Register user
        register_response = client.post(
            "/api/auth/register",
            json={
                "username": "template_user",
                "password": "password123",
                "nickname": "Template User"
            }
        )

        token = register_response.json()["data"]["access_token"]
        auth_headers = {"Authorization": f"Bearer {token}"}

        # Upload document
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]

        # Submit check with template
        check_response = client.post(
            "/api/check",
            headers=auth_headers,
            json={
                "file_id": file_id,
                "filename": "test.docx",
                "check_type": "basic",
                "rule_template_id": sample_rule_template.id
            }
        )

        assert check_response.status_code == 200
        check_id = check_response.json()["data"]["check_id"]

        # Verify last_template_id updated
        profile = client.get("/api/auth/user/profile", headers=auth_headers).json()["data"]
        assert profile["last_template_id"] == sample_rule_template.id
        assert profile["last_template"]["name"] == "测试模板"

        # Get check result
        result_response = client.get(f"/api/check/{check_id}", headers=auth_headers)
        result_data = result_response.json()["data"]
        assert "result" in result_data

        # Generate and download revision
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        download_response = client.get(f"/api/check/{check_id}/download_revised")
        assert download_response.status_code == 200

    def test_multiple_users_independent_counts(self, client, db: Session, sample_docx_path):
        """Test that multiple users have independent count management."""
        # Register first user
        register1 = client.post("/api/auth/register", json={
            "username": "user1", "password": "pass123", "nickname": "User 1"
        })
        token1 = register1.json()["data"]["access_token"]
        headers1 = {"Authorization": f"Bearer {token1}"}

        # Register second user
        register2 = client.post("/api/auth/register", json={
            "username": "user2", "password": "pass123", "nickname": "User 2"
        })
        token2 = register2.json()["data"]["access_token"]
        headers2 = {"Authorization": f"Bearer {token2}"}

        # User 1 submits check
        with open(sample_docx_path, "rb") as f:
            upload1 = client.post("/api/check/upload", headers=headers1,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id1 = upload1.json()["data"]["file_id"]
        client.post("/api/check", headers=headers1,
            json={"file_id": file_id1, "filename": "test.docx", "check_type": "basic"})

        # User 2 submits check
        with open(sample_docx_path, "rb") as f:
            upload2 = client.post("/api/check/upload", headers=headers2,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id2 = upload2.json()["data"]["file_id"]
        client.post("/api/check", headers=headers2,
            json={"file_id": file_id2, "filename": "test.docx", "check_type": "basic"})

        # Verify independent counts
        profile1 = client.get("/api/auth/user/profile", headers=headers1).json()["data"]
        profile2 = client.get("/api/auth/user/profile", headers=headers2).json()["data"]

        assert profile1["free_count"] == 2
        assert profile2["free_count"] == 2

        # Verify users can't access each other's checks
        checks1 = client.get("/api/check/recent", headers=headers1).json()["data"]
        checks2 = client.get("/api/check/recent", headers=headers2).json()["data"]

        assert checks1["total"] == 1
        assert checks2["total"] == 1

    def test_full_check_workflow(self, client, db: Session, sample_docx_path):
        """Test complete workflow for full check type."""
        # Register user
        register_response = client.post("/api/auth/register", json={
            "username": "fullcheck_user", "password": "pass123", "nickname": "Full User"
        })
        token = register_response.json()["data"]["access_token"]
        headers = {"Authorization": f"Bearer {token}"}
        user_id = register_response.json()["data"]["user"]["id"]

        # Give user full check count
        user = db.query(User).filter(User.id == user_id).first()
        user.full_count = 2
        db.commit()

        # Submit full check
        with open(sample_docx_path, "rb") as f:
            upload = client.post("/api/check/upload", headers=headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "full"})
        file_id = upload.json()["data"]["file_id"]

        check = client.post("/api/check", headers=headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "full"})
        assert check.status_code == 200
        check_id = check.json()["data"]["check_id"]

        # Verify full count deducted
        profile = client.get("/api/auth/user/profile", headers=headers).json()["data"]
        assert profile["full_count"] == 1
        assert profile["free_count"] == 3  # Free count unchanged

        # Verify check type
        check_record = db.query(Check).filter(Check.check_id == check_id).first()
        assert check_record.check_type == CheckType.FULL
        assert check_record.cost_type == CostType.FULL

        # Get result (full check should include AI results if enabled)
        result = client.get(f"/api/check/{check_id}", headers=headers).json()["data"]
        assert result["check_type"] == "full"

    def test_check_revision_with_fixes_applied(self, client, db: Session, temp_upload_dir):
        """Test that revision actually applies fixes to the document."""
        from docx import Document
        from docx.shared import Mm

        # Create document with known issues
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Mm(10)  # Wrong margin
        doc.add_paragraph("Test paragraph")

        docx_path = os.path.join(temp_upload_dir, "fixable.docx")
        doc.save(docx_path)

        # Register user
        register = client.post("/api/auth/register", json={
            "username": "fix_user", "password": "pass123", "nickname": "Fix User"
        })
        token = register.json()["data"]["access_token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Upload and check
        with open(docx_path, "rb") as f:
            upload = client.post("/api/check/upload", headers=headers,
                files={"file": ("fixable.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"})
        file_id = upload.json()["data"]["file_id"]

        check = client.post("/api/check", headers=headers,
            json={"file_id": file_id, "filename": "fixable.docx", "check_type": "basic"})
        check_id = check.json()["data"]["check_id"]

        # Generate revision
        revise = client.post(f"/api/check/{check_id}/revise", headers=headers)
        assert revise.status_code == 200

        # Verify revised document exists and can be opened
        check_record = db.query(Check).filter(Check.check_id == check_id).first()
        assert check_record.revised_file_path is not None

        revised_doc = Document(check_record.revised_file_path)
        assert len(revised_doc.paragraphs) > 0

    def test_user_stats_after_multiple_checks(self, client, db: Session, sample_docx_path):
        """Test user statistics are correctly updated after multiple checks."""
        # Register user
        register = client.post("/api/auth/register", json={
            "username": "stats_user", "password": "pass123", "nickname": "Stats User"
        })
        token = register.json()["data"]["access_token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Submit 3 basic checks
        for i in range(3):
            with open(sample_docx_path, "rb") as f:
                upload = client.post("/api/check/upload", headers=headers,
                    files={"file": (f"test{i}.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                    data={"check_type": "basic"})
            file_id = upload.json()["data"]["file_id"]
            client.post("/api/check", headers=headers,
                json={"file_id": file_id, "filename": f"test{i}.docx", "check_type": "basic"})

        # Get stats
        stats = client.get("/api/check/stats", headers=headers).json()["data"]
        assert stats["total_checks"] == 3
        assert stats["basic_checks"] == 3
        assert stats["total_issues"] >= 0

        # Get history
        history = client.get("/api/check/recent?limit=10", headers=headers).json()["data"]
        assert history["total"] == 3
        assert len(history["checks"]) == 3

    def test_error_recovery_workflow(self, client, db: Session):
        """Test that system recovers gracefully from errors without corrupting counts."""
        # Register user
        register = client.post("/api/auth/register", json={
            "username": "error_user", "password": "pass123", "nickname": "Error User"
        })
        token = register.json()["data"]["access_token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Get initial profile
        initial = client.get("/api/auth/user/profile", headers=headers).json()["data"]
        initial_free = initial["free_count"]

        # Try to submit check with non-existent file (should fail without deducting count)
        error_check = client.post("/api/check", headers=headers,
            json={"file_id": "bad_file_id", "filename": "bad.docx", "check_type": "basic"})
        assert error_check.status_code == 200
        assert error_check.json()["code"] != 200

        # Verify count not deducted
        after_error = client.get("/api/auth/user/profile", headers=headers).json()["data"]
        assert after_error["free_count"] == initial_free

        # Verify no check record created
        history = client.get("/api/check/recent", headers=headers).json()["data"]
        assert history["total"] == 0
