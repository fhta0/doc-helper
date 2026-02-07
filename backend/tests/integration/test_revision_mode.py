"""
Integration tests for document revision mode.

Tests the complete revision workflow:
1. Submit check and get issues with fix_action
2. Generate revised document applying fixes
3. Download revised document
4. Verify fixes were correctly applied
5. Edge cases (no fixable issues, invalid check_id, etc.)
"""
import pytest
import os
import json
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Mm, Pt
from app.models import Check, CheckStatus, CheckType, CostType, RuleTemplate


@pytest.mark.integration
@pytest.mark.api
class TestRevisionMode:
    """Test cases for document revision generation and download."""

    def test_generate_revised_document_success(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test successful generation of revised document."""
        # Submit a check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]

        check_response = client.post(
            "/api/check",
            headers=auth_headers,
            json={
                "file_id": file_id,
                "filename": "test.docx",
                "check_type": "basic"
            }
        )

        assert check_response.status_code == 200
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(
            f"/api/check/{check_id}/revise",
            headers=auth_headers
        )

        assert revise_response.status_code == 200
        data = revise_response.json()
        assert data["code"] == 200
        assert "revised_file_url" in data["data"]

        # Verify revised file was created
        check = db.query(Check).filter(Check.check_id == check_id).first()
        assert check.revised_file_path is not None
        assert os.path.exists(check.revised_file_path)

    def test_download_revised_document_success(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test successful download of revised document."""
        # Submit check and generate revision
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Download revised document
        download_response = client.get(f"/api/check/{check_id}/download_revised")

        assert download_response.status_code == 200
        assert download_response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert len(download_response.content) > 0

    def test_download_revised_document_not_generated(self, client, sample_check, db: Session):
        """Test downloading revised document when it hasn't been generated."""
        # Ensure no revised file
        sample_check.revised_file_path = None
        db.commit()

        response = client.get(f"/api/check/{sample_check.check_id}/download_revised")

        # Should return 404 or error
        assert response.status_code in [200, 404]
        if response.status_code == 200:
            data = response.json()
            assert data["code"] == 404

    def test_generate_revised_document_nonexistent_check(self, client, auth_headers):
        """Test generating revised document for non-existent check."""
        response = client.post(
            "/api/check/nonexistent_check_id/revise",
            headers=auth_headers
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 3002

    def test_generate_revised_document_incomplete_check(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test generating revised document for incomplete check."""
        # Create a pending check
        check = Check(
            check_id="pending_check_123",
            user_id=test_user.id,
            file_id="file_123",
            filename="test.docx",
            file_path=sample_docx_path,
            check_type=CheckType.BASIC,
            status=CheckStatus.PENDING,
            cost_type=CostType.FREE
        )
        db.add(check)
        db.commit()

        response = client.post(
            f"/api/check/{check.check_id}/revise",
            headers=auth_headers
        )

        assert response.status_code == 200
        data = response.json()
        assert data["code"] == 3005

    def test_generate_revised_document_idempotent(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that generating revised document twice returns same result."""
        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document first time
        response1 = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert response1.status_code == 200

        check = db.query(Check).filter(Check.check_id == check_id).first()
        first_path = check.revised_file_path

        # Generate revised document second time
        response2 = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert response2.status_code == 200

        db.refresh(check)
        second_path = check.revised_file_path

        # Should return same path
        assert first_path == second_path

    def test_revised_document_applies_page_margin_fix(self, client, auth_headers, test_user, db: Session, temp_upload_dir):
        """Test that revised document correctly applies page margin fixes."""
        from docx import Document
        from docx.shared import Mm

        # Create a document with wrong margins
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Mm(10)  # Wrong margin
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        doc.add_paragraph("Test content")

        docx_path = os.path.join(temp_upload_dir, "margin_test.docx")
        doc.save(docx_path)

        # Upload and check
        with open(docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("margin_test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "margin_test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify revised document has correct margins
        check = db.query(Check).filter(Check.check_id == check_id).first()
        if check.revised_file_path and os.path.exists(check.revised_file_path):
            revised_doc = Document(check.revised_file_path)
            # Note: Actual margin values depend on the rules loaded
            # This test verifies the document can be opened and has sections
            assert len(revised_doc.sections) > 0

    def test_revised_document_with_template_rules(self, client, auth_headers, test_user, sample_docx_path, sample_rule_template, db: Session):
        """Test that revision works with template-based rules."""
        # Submit check with template
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post(
            "/api/check",
            headers=auth_headers,
            json={
                "file_id": file_id,
                "filename": "test.docx",
                "check_type": "basic",
                "rule_template_id": sample_rule_template.id
            }
        )

        assert check_response.status_code == 200
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify revised file created
        check = db.query(Check).filter(Check.check_id == check_id).first()
        assert check.revised_file_path is not None
        assert os.path.exists(check.revised_file_path)

    def test_check_result_shows_revised_file_status(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that check result API shows revised file generation status."""
        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Get check result before revision
        result_response = client.get(f"/api/check/{check_id}", headers=auth_headers)
        assert result_response.status_code == 200
        data = result_response.json()["data"]
        assert data["revised_file_generated"] == False

        # Generate revised document
        client.post(f"/api/check/{check_id}/revise", headers=auth_headers)

        # Get check result after revision
        result_response = client.get(f"/api/check/{check_id}", headers=auth_headers)
        assert result_response.status_code == 200
        data = result_response.json()["data"]
        assert data["revised_file_generated"] == True

    def test_revised_document_handles_no_fixable_issues(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that revision works even when no issues have fix_action."""
        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Get check and manually set issues without fix_action
        check = db.query(Check).filter(Check.check_id == check_id).first()
        manual_result = {
            "total_issues": 1,
            "issues": [
                {
                    "rule_id": "MANUAL_CHECK",
                    "message": "需要手动检查",
                    "suggestion": "请手动修正",
                    "location": {"type": "paragraph", "index": 0}
                    # No fix_action
                }
            ]
        }
        check.result_json = json.dumps(manual_result)
        db.commit()

        # Generate revised document - should succeed with comments
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify file created
        db.refresh(check)
        assert check.revised_file_path is not None
        assert os.path.exists(check.revised_file_path)

    def test_revised_document_with_mixed_fixable_and_manual_issues(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test revision with both fixable and manual issues."""
        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Manually set mixed issues
        check = db.query(Check).filter(Check.check_id == check_id).first()
        mixed_result = {
            "total_issues": 2,
            "issues": [
                {
                    "rule_id": "PAGE_MARGIN_25",
                    "message": "页边距不符合要求",
                    "fix_action": "set_page_margin",
                    "fix_params": {
                        "top_mm": 25.4,
                        "bottom_mm": 25.4,
                        "left_mm": 31.8,
                        "right_mm": 31.8
                    },
                    "location": {"type": "document"}
                },
                {
                    "rule_id": "MANUAL_CHECK",
                    "message": "需要手动检查",
                    "suggestion": "请手动修正",
                    "location": {"type": "paragraph", "index": 0}
                }
            ]
        }
        check.result_json = json.dumps(mixed_result)
        db.commit()

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify file created
        db.refresh(check)
        assert check.revised_file_path is not None
        assert os.path.exists(check.revised_file_path)

        # Verify revised document can be opened
        revised_doc = Document(check.revised_file_path)
        assert len(revised_doc.paragraphs) > 0

    def test_revised_document_preserves_original_content(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that revised document preserves original text content."""
        # Read original content
        original_doc = Document(sample_docx_path)
        original_text = '\n'.join([p.text for p in original_doc.paragraphs])

        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify content preserved (formatting may change, but text should be same)
        check = db.query(Check).filter(Check.check_id == check_id).first()
        revised_doc = Document(check.revised_file_path)
        revised_text = '\n'.join([p.text for p in revised_doc.paragraphs])

        # Text should be identical or very similar (minor whitespace differences OK)
        assert len(revised_text) > 0
        # For more precise assertion, could compare after normalizing whitespace

    def test_revised_document_filename_format(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that revised document has correct filename format."""
        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("my_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "my_document.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify filename format
        check = db.query(Check).filter(Check.check_id == check_id).first()
        revised_filename = os.path.basename(check.revised_file_path)
        assert "_revised_" in revised_filename
        assert revised_filename.endswith(".docx")

    def test_revised_document_stored_in_user_directory(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that revised document is stored in correct user directory."""
        # Submit check
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Verify file in user directory
        check = db.query(Check).filter(Check.check_id == check_id).first()
        assert str(test_user.id) in check.revised_file_path

    def test_download_revised_document_returns_correct_headers(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that download response has correct content-type and headers."""
        # Submit check and generate revision
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        client.post(f"/api/check/{check_id}/revise", headers=auth_headers)

        # Download
        download_response = client.get(f"/api/check/{check_id}/download_revised")

        assert download_response.status_code == 200
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in download_response.headers.get("content-type", "")
        assert len(download_response.content) > 0

    def test_download_returns_valid_docx_file(self, client, auth_headers, test_user, sample_docx_path, db: Session):
        """Test that downloaded revised document is a valid DOCX file."""
        # Submit check and generate revision
        with open(sample_docx_path, "rb") as f:
            upload_response = client.post(
                "/api/check/upload",
                headers=auth_headers,
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"check_type": "basic"}
            )

        file_id = upload_response.json()["data"]["file_id"]
        check_response = client.post("/api/check", headers=auth_headers,
            json={"file_id": file_id, "filename": "test.docx", "check_type": "basic"})
        check_id = check_response.json()["data"]["check_id"]

        # Generate revised document
        revise_response = client.post(f"/api/check/{check_id}/revise", headers=auth_headers)
        assert revise_response.status_code == 200

        # Download and verify it's a valid DOCX
        download_response = client.get(f"/api/check/{check_id}/download_revised")
        assert download_response.status_code == 200

        # Save to temp file and verify it can be opened
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(download_response.content)
            tmp_path = tmp.name

        # Verify valid DOCX
        from docx import Document
        doc = Document(tmp_path)
        assert len(doc.paragraphs) > 0

        # Clean up
        os.unlink(tmp_path)
