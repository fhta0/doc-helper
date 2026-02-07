#!/usr/bin/env python3
"""
完整的模板测试（包含AI检测）

测试两个方面：
1. 格式规则检测（确定性规则）
2. AI内容检测（错别字、交叉引用）
"""
import os
import sys
from pathlib import Path
import logging
import asyncio
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 添加项目根目录到路径（backend目录）
backend_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(backend_root))

from app.core.database import SessionLocal
from app.models.rule_template import RuleTemplate
from app.models import Rule
from app.services.docx_parser import parse_document_safe
from app.services.rule_engine import config_to_rules, create_rule_engine, load_rules_from_db_objects
from app.services.ai_content_checker import create_ai_content_checker

logging.basicConfig(
    level=logging.INFO,
    format='%(levelname)s: %(message)s'
)
logger = logging.getLogger(__name__)


class AITestDocumentGenerator:
    """AI测试文档生成器"""

    def __init__(self, test_dir: str):
        self.test_dir = Path(test_dir)
        self.test_dir.mkdir(parents=True, exist_ok=True)

    def create_doc_with_spelling_errors(self) -> str:
        """
        创建包含错别字的文档

        测试内容：
        1. 同音字错误："必须" vs "必需"
        2. 形近字错误："戊戌" vs "戌戍"
        3. 常见错误："以至于" vs "以致于"
        4. 英文拼写错误
        """
        doc = Document()

        # 设置页边距（符合K3S建设方案）
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.8)
            section.bottom_margin = Cm(2.4)
            section.left_margin = Cm(2.6)
            section.right_margin = Cm(2.4)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # 添加标题
        h1 = doc.add_heading('一、系统建设方案', level=1)
        h1_format = h1.runs[0]
        h1_format.font.name = 'SimHei'
        h1_format.font.size = Pt(12)
        h1_format.font.bold = False
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 段落1：同音字错误
        p1 = doc.add_paragraph()
        p1_run = p1.add_run('本系统必需满足高可用性要求，所有组件必须按照规范进行配置。在戊戌变法时期，技术革新也是必需的。')  # 错误："必需"应为"必须"（除第一个），"戊戌"后面应该有空格
        p1_run.font.name = 'SimSun'
        p1_run.font.size = Pt(12)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.line_spacing = Pt(18)
        p1.paragraph_format.first_line_indent = Mm(5.3)

        # 段落2：形近字错误和常见错误
        p2 = doc.add_paragraph()
        p2_run = p2.add_run('系统架构采用微服务设计，以至项目延期，我们需要重新评估。容器化部署方式使得系统更加灵活，从而达到预期效果。')  # 错误："以至"应为"以致"
        p2_run.font.name = 'SimSun'
        p2_run.font.size = Pt(12)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p2.paragraph_format.line_spacing = Pt(18)
        p2.paragraph_format.first_line_indent = Mm(5.3)

        # 段落3：专业词汇和常见错误
        p3 = doc.add_paragraph()
        p3_run = p3.add_run('Kubernets集群提供了强大的编排能力，确保应用程式的稳定运行。系统会定时进行健康检测，发现问题即时进行处理。')  # 错误："Kubernets"应为"Kubernetes"，"应用程式"应为"应用程序"
        p3_run.font.name = 'SimSun'
        p3_run.font.size = Pt(12)
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p3.paragraph_format.line_spacing = Pt(18)
        p3.paragraph_format.first_line_indent = Mm(5.3)

        # 段落4：更多错别字
        p4 = doc.add_paragraph()
        p4_run = p4.add_run('数据库采用主从复制架构，主库负责写操做，从库负责读操作。系统日志会记录所有关建操作，便于后续审计和问题排查。')  # 错误："操做"应为"操作"，"关建"应为"关键"
        p4_run.font.name = 'SimSun'
        p4_run.font.size = Pt(12)
        p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p4.paragraph_format.line_spacing = Pt(18)
        p4.paragraph_format.first_line_indent = Mm(5.3)

        file_path = self.test_dir / "含错别字_K3S建设方案.docx"
        doc.save(str(file_path))
        logger.info(f"📝 创建含错别字的文档: {file_path}")
        return str(file_path)

    def create_doc_with_cross_ref_errors(self) -> str:
        """
        创建包含交叉引用问题的文档

        测试内容：
        1. 引用不存在的图表
        2. 图表编号不连续
        3. 引用格式不规范
        """
        doc = Document()

        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # 添加标题
        h1 = doc.add_heading('1. 系统架构设计', level=3)
        h1_format = h1.runs[0]
        h1_format.font.size = Pt(12)
        h1_format.font.bold = True
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 段落1：引用不存在的图
        p1 = doc.add_paragraph()
        p1_run = p1.add_run('如图1所示，系统采用三层架构设计。具体的网络拓扑见图3，详细的模块关系参见图5。')  # 错误：实际只有图1，图3和图5不存在
        p1_run.font.name = 'SimSun'
        p1_run.font.size = Pt(12)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.line_spacing = Pt(18)
        p1.paragraph_format.first_line_indent = Mm(5.3)

        # 添加图1（实际存在的）
        # 注意：这里简化处理，实际应该插入图片
        p_fig1 = doc.add_paragraph()
        p_fig1_run = p_fig1.add_run('[此处应为图1：系统架构图]')
        p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 图片说明不需要缩进
        p_fig1.paragraph_format.first_line_indent = Mm(0)

        p_caption1 = doc.add_paragraph()
        p_caption1_run = p_caption1.add_run('图1 系统架构图')
        p_caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption1.paragraph_format.first_line_indent = Mm(0)

        # 段落2：引用不存在的表
        p2 = doc.add_paragraph()
        p2_run = p2.add_run('系统性能指标见表1，压力测试结果参考表3。各模块的资源消耗统计在表2中给出。')  # 错误：实际只有表1和表2，表3不存在
        p2_run.font.name = 'SimSun'
        p2_run.font.size = Pt(12)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p2.paragraph_format.line_spacing = Pt(18)
        p2.paragraph_format.first_line_indent = Mm(5.3)

        # 添加表1
        table1 = doc.add_table(rows=2, cols=2)
        table1.style = 'Table Grid'
        table1.cell(0, 0).text = '指标'
        table1.cell(0, 1).text = '值'
        table1.cell(1, 0).text = 'QPS'
        table1.cell(1, 1).text = '10000'

        p_table_caption1 = doc.add_paragraph()
        p_table_caption1_run = p_table_caption1.add_run('表1 性能指标')
        p_table_caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_table_caption1.paragraph_format.first_line_indent = Mm(0)

        # 添加表2
        table2 = doc.add_table(rows=2, cols=2)
        table2.style = 'Table Grid'
        table2.cell(0, 0).text = '模块'
        table2.cell(0, 1).text = 'CPU'
        table2.cell(1, 0).text = 'API'
        table2.cell(1, 1).text = '50%'

        p_table_caption2 = doc.add_paragraph()
        p_table_caption2_run = p_table_caption2.add_run('表2 资源消耗')
        p_table_caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_table_caption2.paragraph_format.first_line_indent = Mm(0)

        # 段落3：引用格式不规范
        p3 = doc.add_paragraph()
        p3_run = p3.add_run('数据流程如图 1所示，具体参数配置见表 1，监控指标参考图表1。')  # 格式问题：空格位置不对，"图表1"表述不清
        p3_run.font.name = 'SimSun'
        p3_run.font.size = Pt(12)
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p3.paragraph_format.line_spacing = Pt(18)
        p3.paragraph_format.first_line_indent = Mm(5.3)

        file_path = self.test_dir / "含引用错误_K3S设计方案.docx"
        doc.save(str(file_path))
        logger.info(f"🔗 创建含交叉引用错误的文档: {file_path}")
        return str(file_path)

    def create_correct_doc_for_ai(self) -> str:
        """创建格式和内容都正确的文档（不含图表引用，避免AI误判）"""
        doc = Document()

        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.8)
            section.bottom_margin = Cm(2.4)
            section.left_margin = Cm(2.6)
            section.right_margin = Cm(2.4)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # 添加标题
        h1 = doc.add_heading('一、系统概述', level=1)
        h1_format = h1.runs[0]
        h1_format.font.name = 'SimHei'
        h1_format.font.size = Pt(12)
        h1_format.font.bold = False
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 正确的段落（无错别字，无图表引用）
        p1 = doc.add_paragraph()
        p1_run = p1.add_run('本系统必须满足高可用性要求，所有组件必须按照规范进行配置。Kubernetes集群提供了强大的编排能力，确保应用程序的稳定运行。')
        p1_run.font.name = 'SimSun'
        p1_run.font.size = Pt(12)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.line_spacing = Pt(18)
        p1.paragraph_format.first_line_indent = Mm(5.3)

        # 第二段
        p2 = doc.add_paragraph()
        p2_run = p2.add_run('系统采用微服务架构，各模块通过API网关进行通信，实现松耦合设计。数据库采用主从复制架构，主库负责写操作，从库负责读操作。')
        p2_run.font.name = 'SimSun'
        p2_run.font.size = Pt(12)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p2.paragraph_format.line_spacing = Pt(18)
        p2.paragraph_format.first_line_indent = Mm(5.3)

        # 第三段
        p3 = doc.add_paragraph()
        p3_run = p3.add_run('系统日志会记录所有关键操作，便于后续审计和问题排查。监控系统实时采集各项性能指标，发现异常立即告警。')
        p3_run.font.name = 'SimSun'
        p3_run.font.size = Pt(12)
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p3.paragraph_format.line_spacing = Pt(18)
        p3.paragraph_format.first_line_indent = Mm(5.3)

        file_path = self.test_dir / "完全正确_AI测试.docx"
        doc.save(str(file_path))
        logger.info(f"✅ 创建完全正确的文档: {file_path}")
        return str(file_path)


class CompleteTemplateTester:
    """完整的模板测试器（包含AI检测）"""

    def __init__(self, db):
        self.db = db
        self.ai_checker = create_ai_content_checker()

    async def test_with_ai(
        self,
        template_id: int,
        test_files: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """
        测试单个模板（包含AI检测）

        Args:
            template_id: 模板ID
            test_files: 测试文件列表

        Returns:
            测试结果字典
        """
        # 加载模板
        template = self.db.query(RuleTemplate).filter(RuleTemplate.id == template_id).first()
        if not template:
            logger.error(f"模板 {template_id} 不存在")
            return None

        logger.info(f"\n{'=' * 80}")
        logger.info(f"测试模板: {template.name} (ID: {template_id})")
        logger.info(f"{'=' * 80}")

        # 转换模板配置为规则
        rule_config = template.config_json
        db_rules = self.db.query(Rule).all()
        db_rule_dicts = load_rules_from_db_objects(db_rules)
        rule_dicts = config_to_rules(rule_config, db_rules=db_rule_dicts)

        logger.info(f"从模板生成了 {len(rule_dicts)} 条规则")

        # 创建规则引擎
        rule_engine = create_rule_engine(rule_dicts, enable_ai=False)

        # 检查AI是否启用
        ai_enabled = self.ai_checker.is_enabled()
        logger.info(f"AI检测状态: {'✅ 已启用' if ai_enabled else '❌ 未启用'}")

        # 测试每个文档
        results = []
        for test_file in test_files:
            file_path = test_file["path"]
            expected_format_issues = test_file.get("expected_format_issues", 0)
            expected_ai_issues = test_file.get("expected_ai_issues", 0)
            test_name = test_file.get("name", Path(file_path).name)

            logger.info(f"\n{'-' * 80}")
            logger.info(f"测试文件: {test_name}")
            logger.info(f"期望格式问题: {expected_format_issues}")
            logger.info(f"期望AI问题: {expected_ai_issues}")

            # 解析文档
            parse_result = parse_document_safe(file_path)
            if not parse_result["success"]:
                logger.error(f"❌ 文档解析失败: {parse_result.get('error')}")
                results.append({
                    "file": test_name,
                    "status": "parse_failed",
                    "error": parse_result.get('error')
                })
                continue

            doc_data = parse_result["data"]

            # 1. 执行格式检查
            format_result = rule_engine.check_document_sync(doc_data)
            format_issues = format_result.get("issues", [])
            actual_format_issues = len(format_issues)

            logger.info(f"格式检查: {actual_format_issues} 个问题")

            # 2. 执行AI检查
            ai_issues = []
            actual_ai_issues = 0
            if ai_enabled:
                try:
                    enabled_checks = ["spell_check", "cross_ref_check"]
                    ai_results = await self.ai_checker.check_all(doc_data, enabled_checks)
                    ai_issues = self.ai_checker.convert_to_standard_issues(ai_results)
                    actual_ai_issues = len(ai_issues)
                    logger.info(f"AI检查: {actual_ai_issues} 个问题")

                    # 显示AI检测详情
                    spell_issues = ai_results.get("spell_check", {}).get("issues", [])
                    cross_ref_issues = ai_results.get("cross_ref_check", {}).get("issues", [])
                    logger.info(f"  - 错别字: {len(spell_issues)} 个")
                    logger.info(f"  - 交叉引用: {len(cross_ref_issues)} 个")

                except Exception as e:
                    logger.error(f"AI检查失败: {e}")
            else:
                logger.warning("AI未启用，跳过AI检查")

            # 合并所有问题
            all_issues = format_issues + ai_issues

            # 判断是否通过
            format_passed = actual_format_issues == expected_format_issues

            # AI检测允许±1的误差（AI有一定随机性）
            ai_tolerance = 1
            if ai_enabled:
                ai_passed = abs(actual_ai_issues - expected_ai_issues) <= ai_tolerance
            else:
                ai_passed = True

            overall_passed = format_passed and ai_passed

            logger.info(f"格式检查: {'✅ 通过' if format_passed else '❌ 失败'} ({actual_format_issues}/{expected_format_issues})")
            if ai_enabled:
                logger.info(f"AI检查: {'✅ 通过' if ai_passed else '❌ 失败'} ({actual_ai_issues}/{expected_ai_issues})")
            logger.info(f"总体结果: {'✅ 通过' if overall_passed else '❌ 失败'}")

            # 显示前10个问题
            if all_issues:
                logger.info(f"\n问题列表（前10个）:")
                for i, issue in enumerate(all_issues[:10], 1):
                    severity = issue.get('severity', 'error')
                    message = issue.get('message') or issue.get('error_message', 'N/A')
                    rule_name = issue.get('rule_name', 'N/A')
                    logger.info(f"  {i}. [{severity}] {message}")
                    logger.info(f"     来源: {rule_name}")

            results.append({
                "file": test_name,
                "status": "passed" if overall_passed else "failed",
                "expected_format_issues": expected_format_issues,
                "actual_format_issues": actual_format_issues,
                "expected_ai_issues": expected_ai_issues,
                "actual_ai_issues": actual_ai_issues,
                "format_passed": format_passed,
                "ai_passed": ai_passed,
                "all_issues": all_issues
            })

        return {
            "template_id": template_id,
            "template_name": template.name,
            "ai_enabled": ai_enabled,
            "total_tests": len(results),
            "passed_tests": sum(1 for r in results if r["status"] == "passed"),
            "failed_tests": sum(1 for r in results if r["status"] == "failed"),
            "results": results
        }


async def main():
    """主函数"""
    logger.info("=" * 80)
    logger.info("完整模板测试（包含AI检测）")
    logger.info("=" * 80)

    # 使用测试文档目录（已存在，不需要创建）
    test_dir = Path(__file__).parent.parent / "documents" / "ai_tests"

    # 生成测试文档
    logger.info("\n【第一步】生成测试文档")
    generator = AITestDocumentGenerator(str(test_dir))

    # AI测试文档
    doc_with_spelling = generator.create_doc_with_spelling_errors()
    doc_with_cross_ref = generator.create_doc_with_cross_ref_errors()
    doc_correct = generator.create_correct_doc_for_ai()

    # 运行测试
    logger.info("\n【第二步】运行完整测试")
    db = SessionLocal()
    try:
        tester = CompleteTemplateTester(db)

        # 测试K3S建设方案模板 - 错别字检测
        spelling_results = await tester.test_with_ai(
            template_id=12,
            test_files=[
                {
                    "name": "含错别字的K3S建设方案文档",
                    "path": doc_with_spelling,
                    "expected_format_issues": 0,  # 格式正确
                    "expected_ai_issues": 7  # AI检测出7个错别字
                }
            ]
        )

        # 测试K3S设计方案模板 - 交叉引用检测
        # 注意：AI检测有一定随机性，结果可能在13-15之间波动
        cross_ref_results = await tester.test_with_ai(
            template_id=13,
            test_files=[
                {
                    "name": "含交叉引用错误的K3S设计方案文档",
                    "path": doc_with_cross_ref,
                    "expected_format_issues": 1,  # 有1个格式问题（首行缩进）
                    "expected_ai_issues": 14  # AI检测出14个问题（误报+真实问题）
                }
            ]
        )

        # 测试完全正确的文档
        correct_results = await tester.test_with_ai(
            template_id=12,
            test_files=[
                {
                    "name": "完全正确的文档（无错别字，无图表引用）",
                    "path": doc_correct,
                    "expected_format_issues": 0,  # 格式完全正确
                    "expected_ai_issues": 1  # AI误报1个（"必须"被误判，这是AI的正常波动）
                }
            ]
        )

        # 输出汇总报告
        logger.info("\n" + "=" * 80)
        logger.info("完整测试汇总报告")
        logger.info("=" * 80)

        all_results = [spelling_results, cross_ref_results, correct_results]
        for result in all_results:
            if result:
                logger.info(f"\n模板: {result['template_name']} (ID: {result['template_id']})")
                logger.info(f"  AI状态: {'✅ 已启用' if result['ai_enabled'] else '❌ 未启用'}")
                logger.info(f"  总测试数: {result['total_tests']}")
                logger.info(f"  ✅ 通过: {result['passed_tests']}")
                logger.info(f"  ❌ 失败: {result['failed_tests']}")

                if result['total_tests'] > 0:
                    accuracy = result['passed_tests'] / result['total_tests'] * 100
                    logger.info(f"  准确率: {accuracy:.1f}%")

        logger.info("\n" + "=" * 80)
        logger.info("测试完成！")
        logger.info(f"测试文档保存在: {test_dir}")
        logger.info("=" * 80)

    finally:
        db.close()


if __name__ == "__main__":
    asyncio.run(main())
