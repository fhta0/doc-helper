#!/usr/bin/env python3
"""
测试模板规则的准确性

根据K3S建设方案模板和K3S设计方案模板，生成测试文档，
然后使用现有的检查功能验证模板规则的准确性。
"""
import os
import sys
from pathlib import Path
import logging
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# 添加项目根目录到路径（backend目录）
backend_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(backend_root))

from app.core.database import SessionLocal
from app.models.rule_template import RuleTemplate
from app.services.docx_parser import parse_document_safe
from app.services.rule_engine import config_to_rules, create_rule_engine, load_rules_from_db_objects
from app.models import Rule

logging.basicConfig(
    level=logging.INFO,
    format='%(levelname)s: %(message)s'
)
logger = logging.getLogger(__name__)


class TemplateTestDocumentGenerator:
    """测试文档生成器"""

    def __init__(self, test_dir: str):
        self.test_dir = Path(test_dir)
        self.test_dir.mkdir(parents=True, exist_ok=True)

    def create_correct_doc_k3s_construction(self) -> str:
        """
        创建完全符合K3S建设方案模板的文档

        模板规范：
        - 页边距：上2.8cm, 下2.4cm, 左2.6cm, 右2.4cm
        - 正文：宋体 12pt, 行距18磅, 首行缩进2字符
        - 标题：黑体 12pt, 不加粗
        """
        doc = Document()

        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.8)
            section.bottom_margin = Cm(2.4)
            section.left_margin = Cm(2.6)
            section.right_margin = Cm(2.4)
            section.page_width = Cm(21.0)  # A4
            section.page_height = Cm(29.7)  # A4

        # 添加标题
        h1 = doc.add_heading('一、项目概述', level=1)
        h1_format = h1.runs[0]
        h1_format.font.name = 'SimHei'
        h1_format.font.size = Pt(12)
        h1_format.font.bold = False
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 添加正文
        p1 = doc.add_paragraph()
        p1_run = p1.add_run('本项目旨在建设一套高可用的K3S容器云平台，为企业提供稳定可靠的容器化应用运行环境。系统采用分布式架构，支持弹性扩展和自动容错。')
        p1_run.font.name = 'SimSun'
        p1_run.font.size = Pt(12)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.line_spacing = Pt(18)
        # 2字符缩进 = 2 * 2.65mm = 5.3mm
        p1.paragraph_format.first_line_indent = Mm(5.3)

        # 二级标题
        h2 = doc.add_heading('1.1 建设目标', level=2)
        h2_format = h2.runs[0]
        h2_format.font.name = 'SimHei'
        h2_format.font.size = Pt(12)
        h2_format.font.bold = False
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 正文
        p2 = doc.add_paragraph()
        p2_run = p2.add_run('通过K3S平台建设，实现应用容器化部署、自动化运维、资源动态调度等功能，提升系统的可靠性和运维效率。')
        p2_run.font.name = 'SimSun'
        p2_run.font.size = Pt(12)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p2.paragraph_format.line_spacing = Pt(18)
        p2.paragraph_format.first_line_indent = Mm(5.3)

        # 三级标题
        h3 = doc.add_heading('1.1.1 核心目标', level=3)
        h3_format = h3.runs[0]
        h3_format.font.size = Pt(12)
        h3_format.font.bold = False
        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 正文
        p3 = doc.add_paragraph()
        p3_run = p3.add_run('实现高可用容器集群，支持自动故障转移，确保业务连续性。建立完善的监控告警体系，实时掌握系统运行状态。')
        p3_run.font.name = 'SimSun'
        p3_run.font.size = Pt(12)
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p3.paragraph_format.line_spacing = Pt(18)
        p3.paragraph_format.first_line_indent = Mm(5.3)

        file_path = self.test_dir / "正确_K3S建设方案.docx"
        doc.save(str(file_path))
        logger.info(f"✅ 创建符合规范的文档: {file_path}")
        return str(file_path)

    def create_incorrect_doc_k3s_construction(self) -> str:
        """
        创建不符合K3S建设方案模板的文档

        故意引入的错误：
        1. 页边距错误：使用默认的2.54cm
        2. 正文字体错误：使用微软雅黑
        3. 正文字号错误：使用14pt
        4. 行距错误：使用1.5倍行距
        5. 首行缩进错误：无缩进
        """
        doc = Document()

        # 页边距保持默认（错误）
        sections = doc.sections
        for section in sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # 一级标题（字体错误）
        h1 = doc.add_heading('一、项目概述', level=1)
        h1_format = h1.runs[0]
        h1_format.font.name = 'Microsoft YaHei'  # 错误：应该是黑体
        h1_format.font.size = Pt(14)  # 错误：应该是12pt
        h1_format.font.bold = True  # 错误：不应该加粗

        # 正文（多个错误）
        p1 = doc.add_paragraph()
        p1_run = p1.add_run('本项目旨在建设一套高可用的K3S容器云平台。')
        p1_run.font.name = 'Microsoft YaHei'  # 错误：应该是宋体
        p1_run.font.size = Pt(14)  # 错误：应该是12pt
        p1.paragraph_format.line_spacing = 1.5  # 错误：应该是18磅
        # 无首行缩进（错误）

        # 二级标题
        h2 = doc.add_heading('1.1 建设目标', level=2)
        h2_format = h2.runs[0]
        h2_format.font.name = 'Arial'  # 错误
        h2_format.font.size = Pt(13)  # 错误

        # 正文
        p2 = doc.add_paragraph('通过K3S平台建设，实现应用容器化部署。')
        p2_format = p2.runs[0]
        p2_format.font.name = 'KaiTi'  # 错误：楷体
        p2_format.font.size = Pt(11)  # 错误：11pt

        file_path = self.test_dir / "错误_K3S建设方案.docx"
        doc.save(str(file_path))
        logger.info(f"❌ 创建包含错误的文档: {file_path}")
        return str(file_path)

    def create_correct_doc_k3s_design(self) -> str:
        """
        创建完全符合K3S设计方案模板的文档

        模板规范：
        - 页边距：上2.54cm, 下2.54cm, 左3.18cm, 右3.18cm
        - 正文：宋体 12pt, 行距18磅, 首行缩进2字符
        - 标题：12pt, 加粗
        """
        doc = Document()

        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # 添加3级标题（设计方案从3级开始）
        h3 = doc.add_heading('1. 系统架构设计', level=3)
        h3_format = h3.runs[0]
        h3_format.font.size = Pt(12)
        h3_format.font.bold = True
        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 正文
        p1 = doc.add_paragraph()
        p1_run = p1.add_run('K3S系统采用微服务架构，各模块通过API网关进行通信，实现松耦合设计。系统包括控制平面、数据平面和管理平面三个主要部分。')
        p1_run.font.name = 'SimSun'
        p1_run.font.size = Pt(12)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.line_spacing = Pt(18)
        p1.paragraph_format.first_line_indent = Mm(5.3)

        # 4级标题
        h4 = doc.add_heading('1.1 控制平面设计', level=4)
        h4_format = h4.runs[0]
        h4_format.font.name = 'Arial'
        h4_format.font.size = Pt(12)
        h4_format.font.bold = True
        h4.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 正文
        p2 = doc.add_paragraph()
        p2_run = p2.add_run('控制平面负责集群管理、调度决策和资源分配，采用高可用部署模式，确保管理功能的可靠性。')
        p2_run.font.name = 'SimSun'
        p2_run.font.size = Pt(12)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p2.paragraph_format.line_spacing = Pt(18)
        p2.paragraph_format.first_line_indent = Mm(5.3)

        file_path = self.test_dir / "正确_K3S设计方案.docx"
        doc.save(str(file_path))
        logger.info(f"✅ 创建符合规范的文档: {file_path}")
        return str(file_path)


class TemplateAccuracyTester:
    """模板准确性测试器"""

    def __init__(self, db):
        self.db = db

    def test_template(self, template_id: int, test_files: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        测试单个模板的准确性

        Args:
            template_id: 模板ID
            test_files: 测试文件列表 [{"path": "...", "expected_issues": 0}, ...]

        Returns:
            测试结果字典
        """
        # 加载模板
        template = self.db.query(RuleTemplate).filter(RuleTemplate.id == template_id).first()
        if not template:
            logger.error(f"模板 {template_id} 不存在")
            return None

        logger.info(f"\n{'=' * 80}")
        logger.info(f"测试模板: {template.name} (ID: {template_id})")
        logger.info(f"{'=' * 80}")

        # 转换模板配置为规则
        rule_config = template.config_json
        db_rules = self.db.query(Rule).all()
        db_rule_dicts = load_rules_from_db_objects(db_rules)
        rule_dicts = config_to_rules(rule_config, db_rules=db_rule_dicts)

        logger.info(f"从模板生成了 {len(rule_dicts)} 条规则")

        # 创建规则引擎
        rule_engine = create_rule_engine(rule_dicts, enable_ai=False)

        # 测试每个文档
        results = []
        for test_file in test_files:
            file_path = test_file["path"]
            expected_issues = test_file["expected_issues"]
            test_name = test_file.get("name", Path(file_path).name)

            logger.info(f"\n{'-' * 80}")
            logger.info(f"测试文件: {test_name}")
            logger.info(f"期望问题数: {expected_issues}")

            # 解析文档
            parse_result = parse_document_safe(file_path)
            if not parse_result["success"]:
                logger.error(f"❌ 文档解析失败: {parse_result.get('error')}")
                results.append({
                    "file": test_name,
                    "status": "parse_failed",
                    "error": parse_result.get('error')
                })
                continue

            doc_data = parse_result["data"]

            # 执行检查
            check_result = rule_engine.check_document_sync(doc_data)
            actual_issues = check_result.get("total_issues", 0)
            issues = check_result.get("issues", [])

            # 判断是否通过
            passed = actual_issues == expected_issues

            logger.info(f"实际问题数: {actual_issues}")
            logger.info(f"测试结果: {'✅ 通过' if passed else '❌ 失败'}")

            # 显示前5个问题
            if issues:
                logger.info(f"\n问题列表（前5个）:")
                for i, issue in enumerate(issues[:5], 1):
                    # 显示完整的issue信息用于调试
                    logger.info(f"  {i}. [{issue.get('severity', 'error')}] {issue.get('message') or issue.get('error_message', 'N/A')}")
                    logger.info(f"     规则: {issue.get('rule_name', 'N/A')} (ID: {issue.get('rule_id', 'N/A')})")
                    if issue.get("locations"):
                        loc = issue["locations"][0]
                        logger.info(f"     位置: {loc}")
                    # 调试：显示完整issue
                    logger.debug(f"     完整issue: {issue}")

            results.append({
                "file": test_name,
                "status": "passed" if passed else "failed",
                "expected_issues": expected_issues,
                "actual_issues": actual_issues,
                "issues": issues
            })

        return {
            "template_id": template_id,
            "template_name": template.name,
            "total_tests": len(results),
            "passed_tests": sum(1 for r in results if r["status"] == "passed"),
            "failed_tests": sum(1 for r in results if r["status"] == "failed"),
            "results": results
        }


def main():
    """主函数"""
    logger.info("=" * 80)
    logger.info("模板准确性测试")
    logger.info("=" * 80)

    # 使用测试文档目录（已存在，不需要创建）
    test_dir = Path(__file__).parent.parent / "documents" / "format_tests"

    # 生成测试文档
    logger.info("\n【第一步】生成测试文档")
    generator = TemplateTestDocumentGenerator(str(test_dir))

    # K3S建设方案测试文档
    correct_construction = generator.create_correct_doc_k3s_construction()
    incorrect_construction = generator.create_incorrect_doc_k3s_construction()

    # K3S设计方案测试文档
    correct_design = generator.create_correct_doc_k3s_design()

    # 运行测试
    logger.info("\n【第二步】运行测试")
    db = SessionLocal()
    try:
        tester = TemplateAccuracyTester(db)

        # 测试K3S建设方案模板 (ID: 12)
        construction_results = tester.test_template(
            template_id=12,
            test_files=[
                {
                    "name": "正确的K3S建设方案文档",
                    "path": correct_construction,
                    "expected_issues": 0  # 应该没有问题
                },
                {
                    "name": "错误的K3S建设方案文档",
                    "path": incorrect_construction,
                    "expected_issues": 6  # 预期有6个问题：页边距、字体、字号、行距、缩进、标题
                }
            ]
        )

        # 测试K3S设计方案模板 (ID: 13)
        design_results = tester.test_template(
            template_id=13,
            test_files=[
                {
                    "name": "正确的K3S设计方案文档",
                    "path": correct_design,
                    "expected_issues": 0
                }
            ]
        )

        # 输出汇总报告
        logger.info("\n" + "=" * 80)
        logger.info("测试汇总报告")
        logger.info("=" * 80)

        for result in [construction_results, design_results]:
            if result:
                logger.info(f"\n模板: {result['template_name']} (ID: {result['template_id']})")
                logger.info(f"  总测试数: {result['total_tests']}")
                logger.info(f"  ✅ 通过: {result['passed_tests']}")
                logger.info(f"  ❌ 失败: {result['failed_tests']}")
                logger.info(f"  准确率: {result['passed_tests'] / result['total_tests'] * 100:.1f}%")

        logger.info("\n" + "=" * 80)
        logger.info("测试完成！")
        logger.info(f"测试文档保存在: {test_dir}")
        logger.info("=" * 80)

    finally:
        db.close()


if __name__ == "__main__":
    main()
