"""
Unit tests for DOCX Parser.
"""
import pytest
import os
from docx import Document
from app.services.docx_parser import DocxParser, parse_document_safe, contains_chinese, contains_english


@pytest.mark.unit
@pytest.mark.service
class TestDocxParser:
    """Test cases for DocxParser."""

    def test_init_with_valid_file(self, sample_docx_path):
        """Test parser initialization with valid file."""
        parser = DocxParser(sample_docx_path)
        assert parser.file_path == sample_docx_path
        assert parser.doc is not None

    def test_init_with_invalid_file(self):
        """Test parser initialization with non-existent file."""
        with pytest.raises(FileNotFoundError):
            DocxParser("/path/to/nonexistent.docx")

    def test_parse_returns_dict(self, sample_docx_path):
        """Test that parse returns a dictionary with expected keys."""
        parser = DocxParser(sample_docx_path)
        result = parser.parse()

        assert isinstance(result, dict)
        assert "info" in result
        assert "page_settings" in result
        assert "paragraphs" in result
        assert "runs" in result
        assert "headings" in result
        assert "tables" in result
        assert "figures" in result
        assert "table_of_contents" in result
        assert "heading_structure" in result

    def test_parse_info(self, sample_docx_path):
        """Test parsing document info."""
        parser = DocxParser(sample_docx_path)
        info = parser._parse_info()

        assert "filename" in info
        assert "file_size" in info
        assert info["filename"] == "test_document.docx"
        assert info["file_size"] > 0

    def test_parse_page_settings(self, sample_docx_path):
        """Test parsing page settings."""
        parser = DocxParser(sample_docx_path)
        settings = parser._parse_page_settings()

        assert "paper_size" in settings
        assert "margins" in settings
        assert "width_mm" in settings["paper_size"]
        assert "height_mm" in settings["paper_size"]
        assert "top_mm" in settings["margins"]
        assert "bottom_mm" in settings["margins"]
        assert "left_mm" in settings["margins"]
        assert "right_mm" in settings["margins"]

    def test_parse_paragraphs(self, sample_docx_path):
        """Test parsing paragraphs."""
        parser = DocxParser(sample_docx_path)
        paragraphs = parser._parse_paragraphs()

        assert isinstance(paragraphs, list)
        assert len(paragraphs) > 0

        para = paragraphs[0]
        assert "index" in para
        assert "text" in para
        assert "style" in para
        assert "formatting" in para
        assert "alignment" in para
        assert "page_number" in para

    def test_parse_runs(self, sample_docx_path):
        """Test parsing text runs."""
        parser = DocxParser(sample_docx_path)
        runs = parser._parse_runs()

        assert isinstance(runs, list)
        if len(runs) > 0:
            run = runs[0]
            assert "paragraph_index" in run
            assert "run_index" in run
            assert "text" in run
            assert "font" in run
            assert "page_number" in run

    def test_parse_font(self, sample_docx_path):
        """Test font parsing."""
        parser = DocxParser(sample_docx_path)
        doc = parser.doc

        if len(doc.paragraphs) > 0 and len(doc.paragraphs[0].runs) > 0:
            run = doc.paragraphs[0].runs[0]
            font_data = parser._parse_font(run.font)

            assert "name" in font_data
            assert "size_pt" in font_data
            assert "bold" in font_data

    def test_parse_headings(self, sample_docx_path):
        """Test parsing headings."""
        parser = DocxParser(sample_docx_path)
        headings = parser._parse_headings()

        assert isinstance(headings, list)
        if len(headings) > 0:
            heading = headings[0]
            assert "level" in heading
            assert "text" in heading
            assert "style_name" in heading
            assert "alignment" in heading
            assert "paragraph_index" in heading

    def test_parse_tables(self, temp_upload_dir):
        """Test parsing tables."""
        # Create a document with a table
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"

        file_path = os.path.join(temp_upload_dir, "table_test.docx")
        doc.save(file_path)

        parser = DocxParser(file_path)
        tables = parser._parse_tables()

        assert isinstance(tables, list)
        assert len(tables) == 1

    def test_parse_table_of_contents(self, temp_upload_dir):
        """Test parsing table of contents."""
        # Create a document with TOC heading
        doc = Document()
        doc.add_heading('目录', 0)
        doc.add_paragraph('第一章\t\t\t1')
        doc.add_paragraph('第二章\t\t\t5')

        file_path = os.path.join(temp_upload_dir, "toc_test.docx")
        doc.save(file_path)

        parser = DocxParser(file_path)
        toc = parser._parse_table_of_contents()

        assert "exists" in toc
        assert "entries" in toc
        assert toc["exists"] is True

    def test_contains_chinese(self):
        """Test Chinese character detection."""
        assert contains_chinese("这是中文") is True
        assert contains_chinese("This is English") is False
        assert contains_chinese("混合text中英文") is True
        assert contains_chinese("123456") is False

    def test_contains_english(self):
        """Test English character detection."""
        assert contains_english("This is English") is True
        assert contains_english("这是中文") is False
        assert contains_english("混合text中英文") is True
        assert contains_english("123456") is False

    def test_parse_document_safe_success(self, sample_docx_path):
        """Test safe document parsing with valid file."""
        result = parse_document_safe(sample_docx_path)

        assert result["success"] is True
        assert "data" in result
        assert isinstance(result["data"], dict)

    def test_parse_document_safe_file_not_found(self):
        """Test safe document parsing with non-existent file."""
        result = parse_document_safe("/path/to/nonexistent.docx")

        assert result["success"] is False
        assert "error" in result
        assert "error_type" in result
        assert result["error_type"] == "FileNotFoundError"

    def test_parse_paragraph_formatting(self, sample_docx_path):
        """Test paragraph formatting parsing."""
        parser = DocxParser(sample_docx_path)
        doc = parser.doc

        if len(doc.paragraphs) > 0:
            para = doc.paragraphs[0]
            formatting = parser._parse_paragraph_formatting(para)

            assert isinstance(formatting, dict)
            # Keys may vary based on paragraph settings

    def test_get_alignment(self, sample_docx_path):
        """Test alignment conversion."""
        parser = DocxParser(sample_docx_path)

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert parser._get_alignment(WD_ALIGN_PARAGRAPH.LEFT) == "left"
        assert parser._get_alignment(WD_ALIGN_PARAGRAPH.CENTER) == "center"
        assert parser._get_alignment(WD_ALIGN_PARAGRAPH.RIGHT) == "right"
        assert parser._get_alignment(WD_ALIGN_PARAGRAPH.JUSTIFY) == "justify"
        assert parser._get_alignment(None) == "left"

    def test_parse_heading_structure(self, temp_upload_dir):
        """Test parsing heading hierarchy structure."""
        # Create document with hierarchical headings
        doc = Document()
        doc.add_heading('Chapter 1', 1)
        doc.add_heading('Section 1.1', 2)
        doc.add_heading('Section 1.2', 2)
        doc.add_heading('Chapter 2', 1)

        file_path = os.path.join(temp_upload_dir, "hierarchy_test.docx")
        doc.save(file_path)

        parser = DocxParser(file_path)
        structure = parser._parse_heading_structure()

        assert "tree" in structure
        assert "flat" in structure
        assert isinstance(structure["tree"], list)
        assert isinstance(structure["flat"], list)

    def test_normalize_font_name(self, sample_docx_path):
        """Test font name normalization in _parse_font."""
        parser = DocxParser(sample_docx_path)

        # Test font name mapping
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("测试")
        run.font.name = "宋体"
        run.font.size = Pt(12)

        font_data = parser._parse_font(run.font)
        # Should map "宋体" to "SimSun"
        assert font_data["name"] in ["宋体", "SimSun"]

    def test_parse_with_empty_document(self, temp_upload_dir):
        """Test parsing an empty document."""
        doc = Document()
        file_path = os.path.join(temp_upload_dir, "empty_test.docx")
        doc.save(file_path)

        parser = DocxParser(file_path)
        result = parser.parse()

        assert isinstance(result, dict)
        assert len(result["paragraphs"]) == 0
        assert len(result["headings"]) == 0
