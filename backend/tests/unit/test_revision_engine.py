"""
Unit tests for Revision Engine.
"""
import pytest
import os
import tempfile
import shutil
from docx import Document
from app.services.revision_engine import RevisionEngine


@pytest.mark.unit
@pytest.mark.service
class TestRevisionEngine:
    """Test cases for RevisionEngine."""

    @pytest.fixture
    def temp_dir(self):
        """Create a temporary directory for test files."""
        temp_path = tempfile.mkdtemp()
        yield temp_path
        shutil.rmtree(temp_path, ignore_errors=True)

    @pytest.fixture
    def sample_docx(self, temp_dir):
        """Create a sample DOCX file for testing."""
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is the first paragraph.')
        doc.add_paragraph('This is the second paragraph with 错误字.')

        file_path = os.path.join(temp_dir, 'test.docx')
        doc.save(file_path)
        return file_path

    def test_init(self, sample_docx, temp_dir):
        """Test RevisionEngine initialization."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        assert engine.original_file_path == sample_docx
        assert engine.output_dir == output_dir
        assert os.path.exists(output_dir)

    def test_generate_revised_document_no_issues(self, sample_docx, temp_dir):
        """Test generating revised document with no issues."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = []
        revised_path = engine.generate_revised_document(issues)

        assert os.path.exists(revised_path)
        assert revised_path.startswith(output_dir)
        assert 'revised' in os.path.basename(revised_path)

        # Verify the document can be opened
        doc = Document(revised_path)
        assert len(doc.paragraphs) > 0

    def test_generate_revised_document_with_page_margin_fix(self, sample_docx, temp_dir):
        """Test generating revised document with page margin fix."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "MARGIN_CHECK",
                "fix_action": "set_page_margin",
                "fix_params": {
                    "top_mm": 25.4,
                    "bottom_mm": 25.4,
                    "left_mm": 31.8,
                    "right_mm": 31.8
                },
                "suggestion": "设置页边距",
                "location": {"type": "document"}
            }
        ]

        revised_path = engine.generate_revised_document(issues)

        assert os.path.exists(revised_path)

        # Verify margins were set
        doc = Document(revised_path)
        from docx.shared import Mm
        section = doc.sections[0]
        assert abs(section.top_margin.mm - 25.4) < 1
        assert abs(section.left_margin.mm - 31.8) < 1

    def test_generate_revised_document_with_indent_fix(self, sample_docx, temp_dir):
        """Test generating revised document with paragraph indent fix."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "INDENT_CHECK",
                "fix_action": "set_paragraph_indent",
                "fix_params": {
                    "first_line_indent_chars": 2
                },
                "suggestion": "设置首行缩进",
                "location": {"type": "paragraph", "index": 1}
            }
        ]

        revised_path = engine.generate_revised_document(issues)

        assert os.path.exists(revised_path)

    def test_generate_revised_document_with_text_replacement(self, sample_docx, temp_dir):
        """Test generating revised document with text replacement."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "AI_SPELL_CHECK",
                "fix_action": "replace_text",
                "fix_params": {
                    "original": "错误字",
                    "correction": "正确字",
                    "paragraph_index": 2
                },
                "suggestion": "修正错别字",
                "location": {"type": "paragraph", "index": 2}
            }
        ]

        revised_path = engine.generate_revised_document(issues)

        assert os.path.exists(revised_path)

        # Verify text was replaced (with track changes)
        doc = Document(revised_path)
        # The document should have track changes enabled
        # Note: Actual verification of track changes requires XML inspection

    def test_generate_revised_document_with_heading_style_fix(self, sample_docx, temp_dir):
        """Test generating revised document with heading style fix."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "HEADING_STYLE",
                "fix_action": "set_heading_style",
                "fix_params": {
                    "level1": {
                        "font": "SimHei",
                        "size_pt": 18,
                        "alignment": "center",
                        "bold": True
                    }
                },
                "suggestion": "设置标题格式",
                "location": {"type": "heading", "index": 0, "paragraph_index": 0}
            }
        ]

        revised_path = engine.generate_revised_document(issues)

        assert os.path.exists(revised_path)

    def test_generate_revised_document_with_multiple_locations(self, sample_docx, temp_dir):
        """Test generating revised document with multiple locations for same issue."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "FONT_CHECK",
                "fix_action": "set_run_style",
                "fix_params": {
                    "chinese_font": "SimSun",
                    "size_pt": 14
                },
                "suggestion": "设置字体",
                "raw_locations": [
                    {"type": "paragraph", "index": 1},
                    {"type": "paragraph", "index": 2}
                ]
            }
        ]

        revised_path = engine.generate_revised_document(issues)

        assert os.path.exists(revised_path)

    def test_generate_revised_document_with_invalid_location(self, sample_docx, temp_dir):
        """Test handling of invalid location indices."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "TEST",
                "fix_action": "set_paragraph_indent",
                "fix_params": {"first_line_indent_chars": 2},
                "suggestion": "测试",
                "location": {"type": "paragraph", "index": 999}  # Invalid index
            }
        ]

        # Should not crash
        revised_path = engine.generate_revised_document(issues)
        assert os.path.exists(revised_path)

    def test_generate_revised_document_no_fix_action(self, sample_docx, temp_dir):
        """Test handling of issues without fix_action."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "MANUAL_CHECK",
                "fix_action": None,
                "suggestion": "需要手动检查",
                "location": {"type": "paragraph", "index": 1}
            }
        ]

        # Should not crash, should add comment instead
        revised_path = engine.generate_revised_document(issues)
        assert os.path.exists(revised_path)

    def test_apply_fix_exception_handling(self, sample_docx, temp_dir):
        """Test that fix application exceptions are handled gracefully."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "BAD_FIX",
                "fix_action": "nonexistent_action",
                "fix_params": {},
                "suggestion": "测试异常处理",
                "location": {"type": "paragraph", "index": 1}
            }
        ]

        # Should not crash even with invalid fix action
        revised_path = engine.generate_revised_document(issues)
        assert os.path.exists(revised_path)

    def test_generate_revised_document_filename_format(self, sample_docx, temp_dir):
        """Test that generated filename follows the expected format."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        revised_path = engine.generate_revised_document([])

        filename = os.path.basename(revised_path)
        assert filename.startswith('test_revised_')
        assert filename.endswith('.docx')

    def test_generate_revised_document_preserves_content(self, sample_docx, temp_dir):
        """Test that original content is preserved in revised document."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        original_doc = Document(sample_docx)
        original_text = '\n'.join([p.text for p in original_doc.paragraphs])

        revised_path = engine.generate_revised_document([])

        revised_doc = Document(revised_path)
        revised_text = '\n'.join([p.text for p in revised_doc.paragraphs])

        # Content should be the same when no fixes are applied
        assert original_text == revised_text

    def test_generate_revised_document_with_reference_replacement(self, sample_docx, temp_dir):
        """Test generating revised document with reference replacement."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        # Add a paragraph with a reference
        doc = Document(sample_docx)
        doc.add_paragraph('如图99所示')
        doc.save(sample_docx)

        issues = [
            {
                "rule_id": "AI_CROSS_REF_CHECK",
                "fix_action": "replace_ref",
                "fix_params": {
                    "original_ref": "图99",
                    "suggested_ref": "图1",
                    "paragraph_index": 3
                },
                "suggestion": "修正引用编号",
                "location": {"type": "paragraph", "index": 3}
            }
        ]

        revised_path = engine.generate_revised_document(issues)
        assert os.path.exists(revised_path)

    def test_generate_revised_document_with_set_page_size(self, sample_docx, temp_dir):
        """Test set_page_size action (currently not implemented)."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "PAGE_SIZE",
                "fix_action": "set_page_size",
                "fix_params": {"width_mm": 210, "height_mm": 297},
                "suggestion": "设置页面大小",
                "location": {"type": "document"}
            }
        ]

        # Should not crash even though not implemented
        revised_path = engine.generate_revised_document(issues)
        assert os.path.exists(revised_path)

    def test_generate_revised_document_multiple_issues(self, sample_docx, temp_dir):
        """Test generating revised document with multiple different issues."""
        output_dir = os.path.join(temp_dir, 'output')
        engine = RevisionEngine(sample_docx, output_dir)

        issues = [
            {
                "rule_id": "MARGIN_CHECK",
                "fix_action": "set_page_margin",
                "fix_params": {"top_mm": 25.4, "bottom_mm": 25.4, "left_mm": 31.8, "right_mm": 31.8},
                "suggestion": "设置页边距",
                "location": {"type": "document"}
            },
            {
                "rule_id": "INDENT_CHECK",
                "fix_action": "set_paragraph_indent",
                "fix_params": {"first_line_indent_chars": 2},
                "suggestion": "设置首行缩进",
                "location": {"type": "paragraph", "index": 1}
            },
            {
                "rule_id": "AI_SPELL_CHECK",
                "fix_action": "replace_text",
                "fix_params": {"original": "错误字", "correction": "正确字"},
                "suggestion": "修正错别字",
                "location": {"type": "paragraph", "index": 2}
            }
        ]

        revised_path = engine.generate_revised_document(issues)
        assert os.path.exists(revised_path)

        # Verify document is valid
        doc = Document(revised_path)
        assert len(doc.paragraphs) > 0
